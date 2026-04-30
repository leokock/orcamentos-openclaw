#!/usr/bin/env python3.11
# -*- coding: utf-8 -*-
"""
Gera documento Word com resumo do paramétrico Studios BC
Sem índices detalhados — apenas painel, dados, distribuição e custos
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ===== ESTILOS =====
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

# Margens
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Cores
DARK_BLUE = RGBColor(0x2C, 0x3E, 0x50)
CARTESIAN_ORANGE = RGBColor(0xFF, 0xA5, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
    return h

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def format_brl(val):
    """Formata valor em R$ brasileiro"""
    if val is None:
        return "-"
    if val >= 1_000_000:
        return f"R$ {val:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def format_pct(val):
    return f"{val*100:.1f}%"

def add_table_row(table, cells_data, bold=False, header=False, shade=None):
    row = table.add_row()
    for i, val in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(val)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            if bold or header:
                run.font.bold = True
            if header:
                run.font.color.rgb = WHITE
        if header:
            set_cell_shading(cell, "2C3E50")
        elif shade:
            set_cell_shading(cell, shade)
    return row

# ===== CAPA =====
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ESTUDO PARAMÉTRICO DE CUSTOS")
run.font.size = Pt(24)
run.font.color.rgb = DARK_BLUE
run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Studios — Balneário Camboriú/SC")
run.font.size = Pt(16)
run.font.color.rgb = GRAY

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Fase: Viabilidade")
run.font.size = Pt(12)
run.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

footer_info = [
    "Base de dados: 75 projetos reais calibrados",
    "CUB SC mar/2026: R$ 3.028,45",
    "Data: Março/2026",
]
for line in footer_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

# Rodapé com Cartesian
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Cartesian Engenharia")
run.font.size = Pt(11)
run.font.color.rgb = DARK_BLUE
run.font.bold = True

doc.add_page_break()

# ===== 1. DADOS DO PROJETO =====
add_heading_styled("1. Dados do Projeto", level=1)

p = doc.add_paragraph()
p.add_run("Informações do empreendimento utilizadas como base para o estudo paramétrico.").font.color.rgb = GRAY

table = doc.add_table(rows=0, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

dados = [
    ("Localização", "Balneário Camboriú/SC"),
    ("Tipologia", "Studios — 280 unidades (~30m² médios)"),
    ("Pavimentos", "16 (Térreo + G1 a G3 + Lazer + Tipos)"),
    ("AC Total", "20.280 m²"),
    ("Térreo", "1.716 m² (acesso + garagem)"),
    ("G1 a G3", "5.148 m² (garagem — acima do nível)"),
    ("Lazer + Tipos", "13.416 m²"),
    ("Elevadores", "2"),
    ("Laje", "Cubeta"),
    ("Contenção", "Sem (não há subsolo)"),
    ("Prazo", "36 meses"),
    ("Padrão", "Alto"),
    ("Entrega", "Shell — paredes perimetrais + banheiros (sem layout interno)"),
    ("Fachada", "Mista (tijolo aparente + pele de vidro)"),
    ("CUB Referência", "R$ 3.028,45 (SC mar/2026)"),
]

# Header
add_table_row(table, ["Item", "Valor"], header=True)
for i, (item, val) in enumerate(dados):
    shade = "F5F5F5" if i % 2 == 0 else None
    add_table_row(table, [item, val], shade=shade)

doc.add_paragraph()

# ===== 2. PAINEL EXECUTIVO =====
add_heading_styled("2. Painel Executivo", level=1)

p = doc.add_paragraph()
p.add_run("Indicadores principais do estudo paramétrico.").font.color.rgb = GRAY

# KPIs em tabela limpa
table_kpi = doc.add_table(rows=0, cols=2)
table_kpi.style = 'Table Grid'
table_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER

kpis = [
    ("Custo Total (shell delivery)", "R$ 73.328.012"),
    ("R$/m² Médio AC", "R$ 3.616/m²"),
    ("R$/m² Lazer + Tipos", "R$ 4.556/m²"),
    ("R$/m² Garagem", "R$ 1.779/m²"),
    ("CUB Ratio", "1,19"),
    ("Custo por Unidade", "R$ 261.886/un"),
]

add_table_row(table_kpi, ["Indicador", "Valor"], header=True)
for item, val in kpis:
    add_table_row(table_kpi, [item, val])

doc.add_paragraph()

# Nota shell
p = doc.add_paragraph()
run = p.add_run("Nota: ")
run.font.bold = True
p.add_run("Shell delivery = entrega apenas com paredes do perímetro externo e paredes dos banheiros, sem layout interno dos studios. O cliente finaliza o acabamento conforme necessidade. O custo paramétrico já desconta os itens não executados (pisos, forro, pintura interna, revestimentos, louças/metais e alvenaria de layout).")
p.paragraph_format.space_before = Pt(6)

doc.add_paragraph()

# ===== 3. CUSTOS POR MACROGRUPO =====
add_heading_styled("3. Custos por Macrogrupo", level=1)

p = doc.add_paragraph()
p.add_run("Composição de custos por macrogrupo, com base em 75 projetos reais calibrados para CUB SC mar/2026.").font.color.rgb = GRAY

MACROGRUPOS = [
    "Gerenciamento", "Mov. Terra", "Infraestrutura", "Supraestrutura",
    "Alvenaria", "Impermeabilização", "Instalações", "Sist. Especiais",
    "Climatização", "Rev. Int. Parede", "Teto", "Pisos",
    "Pintura", "Esquadrias", "Louças e Metais", "Fachada",
    "Complementares", "Imprevistos"
]

custos_m2 = {
    "Gerenciamento": 451.85, "Mov. Terra": 17.69, "Infraestrutura": 218.65,
    "Supraestrutura": 747.74, "Alvenaria": 162.97, "Impermeabilização": 61.86,
    "Instalações": 402.02, "Sist. Especiais": 175.04, "Climatização": 56.04,
    "Rev. Int. Parede": 177.03, "Teto": 67.53, "Pisos": 199.91,
    "Pintura": 142.87, "Esquadrias": 467.84, "Louças e Metais": 28.76,
    "Fachada": 221.33, "Complementares": 157.92, "Imprevistos": 55.44
}

total_m2 = sum(custos_m2.values())
ac = 20280

table_custos = doc.add_table(rows=0, cols=4)
table_custos.style = 'Table Grid'
table_custos.alignment = WD_TABLE_ALIGNMENT.CENTER

add_table_row(table_custos, ["Macrogrupo", "R$/m²", "% do Total", "Custo Total (R$)"], header=True)

for mg in MACROGRUPOS:
    v = custos_m2[mg]
    pct = v / total_m2
    total_mg = v * ac
    shade = None
    if pct > 0.10:
        shade = "FFF8E1"  # destaque amarelo pros maiores
    add_table_row(table_custos, [
        mg,
        f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
        f"{pct*100:.1f}%",
        f"R$ {total_mg:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")
    ], shade=shade)

add_table_row(table_custos, [
    "TOTAL",
    f"R$ {total_m2:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
    "100%",
    f"R$ {total_m2 * ac:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")
], bold=True, shade="E8F5E9")

doc.add_paragraph()

# Top 3
p = doc.add_paragraph()
run = p.add_run("Maiores incidências: ")
run.font.bold = True
sorted_mg = sorted(custos_m2.items(), key=lambda x: x[1], reverse=True)[:3]
top3_text = ", ".join([f"{mg} ({v/total_m2*100:.1f}%)" for mg, v in sorted_mg])
p.add_run(top3_text + ".")

doc.add_paragraph()

# ===== 4. DISTRIBUIÇÃO POR TIPO DE ÁREA =====
add_heading_styled("4. Distribuição de Custos por Tipo de Área", level=1)

p = doc.add_paragraph()
p.add_run("Análise separada de custos para garagem e lazer+tipos, considerando a diferença de escopo e acabamento entre essas áreas.").font.color.rgb = GRAY

# Quadro de áreas
add_heading_styled("4.1 Quadro de Áreas", level=2)

table_areas = doc.add_table(rows=0, cols=3)
table_areas.style = 'Table Grid'
table_areas.alignment = WD_TABLE_ALIGNMENT.CENTER

add_table_row(table_areas, ["Setor", "Área (m²)", "% da AC"], header=True)
add_table_row(table_areas, ["Térreo (acesso + garagem)", "1.716", "8,5%"])
add_table_row(table_areas, ["G1 a G3 (garagem)", "5.148", "25,4%"])
add_table_row(table_areas, ["Subtotal Garagem", "6.864", "33,8%"], bold=True, shade="F5F5F5")
add_table_row(table_areas, ["Lazer + Tipos", "13.416", "66,2%"])
add_table_row(table_areas, ["AC Total", "20.280", "100%"], bold=True, shade="E8F5E9")

doc.add_paragraph()

# Resultado distribuição
add_heading_styled("4.2 Custo por Tipo de Área", level=2)

p = doc.add_paragraph()
p.add_run("A garagem tem custo significativamente menor por m² porque não recebe acabamentos, instalações completas, esquadrias ou sistemas especiais na mesma proporção que os pavimentos tipo e lazer.")

doc.add_paragraph()

table_dist = doc.add_table(rows=0, cols=4)
table_dist.style = 'Table Grid'
table_dist.alignment = WD_TABLE_ALIGNMENT.CENTER

add_table_row(table_dist, ["Tipo de Área", "Área (m²)", "R$/m²", "Custo Total (R$)"], header=True)
add_table_row(table_dist, ["Garagem (Térreo + G1-G3)", "6.864", "R$ 1.779", "R$ 12.210.529"], shade="F5F5F5")
add_table_row(table_dist, ["Lazer + Tipos (shell)", "13.416", "R$ 4.556", "R$ 61.117.483"], shade="F5F5F5")
add_table_row(table_dist, ["AC Total", "20.280", "R$ 3.616", "R$ 73.328.012"], bold=True, shade="E8F5E9")

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Ratio Garagem/Tipos: ")
run.font.bold = True
p.add_run("39% — a garagem custa aproximadamente 39% do valor por m² dos pavimentos tipo. Referência de mercado: 50-65% para garagens com subsolo (neste caso sem subsolo, portanto menor).")

doc.add_paragraph()

# ===== 5. AJUSTE SHELL DELIVERY =====
add_heading_styled("4.3 Ajuste Shell Delivery", level=2)

p = doc.add_paragraph()
p.add_run("Como os studios serão entregues sem acabamento interno (apenas paredes perimetrais + banheiro), os seguintes itens foram descontados proporcionalmente:")

doc.add_paragraph()

table_shell = doc.add_table(rows=0, cols=3)
table_shell.style = 'Table Grid'
table_shell.alignment = WD_TABLE_ALIGNMENT.CENTER

add_table_row(table_shell, ["Item", "Economia R$/m² AC", "Justificativa"], header=True)

shell_items = [
    ("Pisos", "R$ 66,24", "Sem pisos nos studios (mantém banheiro)"),
    ("Rev. Int. Parede", "R$ 44,00", "Mantém banheiro + periferia, sem rev. interno apto"),
    ("Pintura", "R$ 35,51", "Sem pintura interna (mantém periferia)"),
    ("Teto/Forro", "R$ 22,38", "Sem forro nos studios (mantém banheiro)"),
    ("Alvenaria", "R$ 20,25", "Sem paredes internas de layout"),
    ("Louças e Metais", "R$ 8,34", "Mantém louças mínimas do banheiro"),
    ("TOTAL ECONOMIA", "R$ 196,71/m² AC", ""),
]

for i, (item, eco, just) in enumerate(shell_items):
    bold = i == len(shell_items) - 1
    shade = "E8F5E9" if bold else None
    add_table_row(table_shell, [item, eco, just], bold=bold, shade=shade)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Sem o desconto shell, o custo dos tipos subiria para aproximadamente R$ 4.850/m² (entrega completa com acabamento interno).")
p.paragraph_format.space_before = Pt(6)

doc.add_paragraph()

# ===== 6. COMPARAÇÃO COM PREMISSA =====
add_heading_styled("5. Comparação com Premissa", level=1)

p = doc.add_paragraph()
p.add_run("Comparação dos valores considerados na viabilidade com o resultado do estudo paramétrico.").font.color.rgb = GRAY

table_comp = doc.add_table(rows=0, cols=4)
table_comp.style = 'Table Grid'
table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER

add_table_row(table_comp, ["Tipo de Área", "Premissa", "Paramétrico", "Diferença"], header=True)
add_table_row(table_comp, ["Lazer + Tipos", "R$ 4.000/m²", "R$ 4.556/m²", "+13,9%"], shade="FFEBEE")
add_table_row(table_comp, ["Garagem", "R$ 2.400/m²", "R$ 1.779/m²", "-25,9%"], shade="FFF8E1")
add_table_row(table_comp, ["Total", "R$ 70.138.000", "R$ 73.328.000", "+4,5%"], bold=True, shade="F5F5F5")

doc.add_paragraph()

# Análise
add_heading_styled("5.1 Análise", level=2)

analise = [
    ("Lazer + Tipos (R$ 4.000/m²):", " Abaixo do paramétrico. Para um empreendimento padrão alto em Balneário Camboriú, com fachada mista (tijolo aparente + pele de vidro), esquadrias de alto desempenho e entrega shell, o paramétrico indica R$ 4.556/m². Com entrega completa, subiria para ~R$ 4.850/m². A premissa de R$ 4.000 está apertada."),
    ("Garagem (R$ 2.400/m²):", " Acima do paramétrico. Para garagem sem subsolo (acima do nível), R$ 2.400/m² está alto. Garagens sem subsolo não têm custos de contenção, impermeabilização pesada ou escavação profunda. O paramétrico indica ~R$ 1.780/m²."),
    ("Total:", " As diferenças se compensam parcialmente. O total fica ~4,5% acima da premissa (R$ 73,3M vs R$ 70,1M), o que é uma diferença razoável para fase de viabilidade. A redistribuição dos valores entre tipos e garagem é mais relevante que a diferença total."),
]

for titulo, texto in analise:
    p = doc.add_paragraph()
    run = p.add_run(titulo)
    run.font.bold = True
    p.add_run(texto)

doc.add_paragraph()

# Sugestão
add_heading_styled("5.2 Sugestão de Valores", level=2)

p = doc.add_paragraph()
p.add_run("Para maior aderência à realidade de mercado, sugerimos os seguintes valores para a viabilidade:")

table_sug = doc.add_table(rows=0, cols=3)
table_sug.style = 'Table Grid'
table_sug.alignment = WD_TABLE_ALIGNMENT.CENTER

add_table_row(table_sug, ["Tipo de Área", "R$/m² Sugerido", "Custo Total"], header=True)
add_table_row(table_sug, ["Lazer + Tipos (shell)", "R$ 4.550/m²", "R$ 61.043.000"], shade="E8F5E9")
add_table_row(table_sug, ["Garagem", "R$ 1.800/m²", "R$ 12.355.000"], shade="E8F5E9")
add_table_row(table_sug, ["TOTAL", "R$ 3.621/m²", "R$ 73.398.000"], bold=True, shade="E8F5E9")

doc.add_paragraph()

# ===== PREMISSAS =====
add_heading_styled("6. Premissas e Limitações", level=1)

premissas = [
    "Estudo paramétrico baseado em 75 projetos reais de obras residenciais verticais em SC",
    "CUB de referência: R$ 3.028,45 (SC, março/2026)",
    "Entrega shell nos studios: apenas paredes perimetrais e banheiros",
    "Garagens G1 a G3 acima do nível (sem subsolo, sem contenção)",
    "Fundação: hélice contínua (premissa padrão para BC)",
    "Laje: cubeta",
    "2 elevadores",
    "Prazo: 36 meses",
    "Fachada: mista (tijolo aparente + pele de vidro)",
    "Valores não incluem terreno, incorporação, marketing ou despesas financeiras",
    "Para orçamento executivo, recomenda-se quantificação detalhada com projetos complementares",
]

for item in premissas:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
doc.add_paragraph()

# Assinatura
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Cartesian Engenharia")
run.font.bold = True
run.font.color.rgb = DARK_BLUE
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Integração de escopo, custo e prazo")
run.font.color.rgb = GRAY
run.font.size = Pt(9)

# Salvar
output_path = os.path.expanduser("~/clawd/orcamento-parametrico/parametricos/studios-bc-estudo-parametrico.docx")
doc.save(output_path)
print(f"✓ Documento gerado: {output_path}")
