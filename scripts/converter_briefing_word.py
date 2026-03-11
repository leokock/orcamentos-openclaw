#!/usr/bin/env python3.11
# -*- coding: utf-8 -*-
"""
Converte briefing Markdown para Word
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def converter_md_para_word(input_md, output_docx):
    """Converte Markdown para Word"""
    
    # Ler o briefing markdown
    with open(input_md, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Criar documento Word
    doc = Document()
    
    # Título
    title = doc.add_heading('BRIEFING PARAMÉTRICO - ARMINIO TAVARES (PLACON)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Processar o conteúdo
    lines = content.split('\n')
    skip_first_title = True
    in_table = False
    table_data = []
    
    for line in lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            continue
        
        # Pular primeiro título (já adicionado)
        if skip_first_title and line_stripped.startswith('# '):
            skip_first_title = False
            continue
        
        # Cabeçalhos
        if line_stripped.startswith('### '):
            doc.add_heading(line_stripped.replace('### ', ''), level=3)
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped.replace('## ', ''), level=2)
        elif line_stripped.startswith('# '):
            doc.add_heading(line_stripped.replace('# ', ''), level=1)
        
        # Linhas de separação
        elif line_stripped.startswith('---'):
            doc.add_paragraph('_' * 80)
        
        # Lista com bullet
        elif line_stripped.startswith('- '):
            text = line_stripped[2:]
            # Remover markdown bold
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        
        # Tabela markdown
        elif '|' in line_stripped:
            if not line_stripped.startswith('|---'):
                # Parse linha da tabela
                cells = [c.strip() for c in line_stripped.split('|')]
                cells = [c for c in cells if c]  # Remove vazios
                
                if not in_table:
                    in_table = True
                    table_data = [cells]
                else:
                    table_data.append(cells)
            else:
                # Linha separadora (ignorar)
                continue
        else:
            # Se estava em tabela, criar a tabela agora
            if in_table:
                if len(table_data) > 0:
                    # Criar tabela
                    num_cols = len(table_data[0])
                    table = doc.add_table(rows=len(table_data), cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for i, row_data in enumerate(table_data):
                        for j, cell_text in enumerate(row_data):
                            if j < num_cols:
                                # Remover markdown
                                cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                                table.rows[i].cells[j].text = cell_text
                                # Negrito na primeira linha
                                if i == 0:
                                    table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
                
                in_table = False
                table_data = []
            
            # Texto normal
            if line_stripped and not line_stripped.startswith('```'):
                # Processar negrito
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line_stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        p.add_run(part[2:-2]).bold = True
                    else:
                        p.add_run(part)
    
    # Salvar
    doc.save(output_docx)
    print(f'✓ Briefing convertido: {output_docx}')

if __name__ == '__main__':
    input_file = 'output/Briefing_Arminio-Tavares_20260309.md'
    output_file = 'output/Briefing_Arminio-Tavares_20260309.docx'
    
    converter_md_para_word(input_file, output_file)
