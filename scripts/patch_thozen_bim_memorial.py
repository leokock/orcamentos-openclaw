#!/usr/bin/env python3
"""Patcha o memorial executivo do Thozen Electra adicionando seção
com quantitativos extraídos do BIM (DXF de AC + Exaustão).

Dados são fixos baseados em dxf-arcondicionado/quantitativos-processados-r05.md
e dxf-exaustao/RESUMO-EXTRACAO.md.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PACOTE = Path.home() / "orcamentos-openclaw" / "base" / "pacotes" / "thozen-electra"
EXEC_DOCX = PACOTE / "executivo-thozen-electra.docx"
PARAM_DOCX = PACOTE / "parametrico-thozen-electra.docx"


def style_h(p, size=14, color="2C3E50"):
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = True
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor.from_string(color)


def add_h(doc, text, size=13):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor.from_string("2C3E50")
    return p


def add_p(doc, text, size=10, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = "Arial"
    r.font.italic = italic
    return p


def add_tab(doc, headers, rows):
    n = len(headers)
    t = doc.add_table(rows=1, cols=n)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor.from_string("FFFFFF")
    for row in rows:
        rr = t.add_row()
        for i, v in enumerate(row[:n]):
            rr.cells[i].text = str(v)
            for p in rr.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Arial"


def patch(docx_path: Path):
    if not docx_path.exists():
        print(f"  SKIP: {docx_path} não existe")
        return False

    doc = Document(str(docx_path))
    doc.add_page_break()
    add_h(doc, "9. Quantitativos extraídos do BIM (DXF)", size=14)
    add_p(doc,
          "Esta seção complementa o orçamento automatizado com quantitativos "
          "extraídos diretamente dos arquivos DXF do projeto (ar-condicionado e "
          "exaustão de churrasqueira). Os números abaixo são literais do projeto, "
          "não vêm da base calibrada.",
          italic=True)
    doc.add_paragraph()

    add_h(doc, "9.1. Sistema de Ar-Condicionado", size=12)
    add_p(doc, "Extraído de dxf-arcondicionado/quantitativos-processados-r05.md", italic=True)

    rows = [
        ("Evaporadoras", "80 un", "1.266.000 BTU/h", "~105,5 TR"),
        ("Condensadoras", "117 un", "1.656.000 BTU/h", "~138,0 TR"),
        ("Total geral", "197 un", "—", "**138 TR (carga térmica total)**"),
    ]
    add_tab(doc, ["Equipamento", "Qtd", "Potência", "TR"], rows)
    add_p(doc, "Distribuição: Térreo, Lazer, Pavimentos Tipo (×24)")
    add_p(doc, "Modelo predominante: Splits 9k/12k/18k/24k/30k BTU/h")
    add_p(doc, "Potência média por evaporadora: 15.825 BTU/h")
    add_p(doc, "Potência média por condensadora: 14.154 BTU/h")
    doc.add_paragraph()

    add_h(doc, "9.2. Sistema de Exaustão de Churrasqueiras", size=12)
    add_p(doc, "Extraído de dxf-exaustao/RESUMO-EXTRACAO.md", italic=True)

    rows = [
        ("Churrasqueiras", "195 un", "—"),
        ("Exaustores", "8 un", "TCV 710 Berliner Luft, 10.600 m³/h, 3,0 kW"),
        ("Inversores de frequência", "8 un", "obrigatórios"),
        ("Dutos galvanizado", "1.400-1.720 m", "chapa #24"),
        ("Prumadas", "8 un", "—"),
    ]
    add_tab(doc, ["Item", "Qtd", "Especificação"], rows)
    add_p(doc, "Estimativa de custo: R$ 1.100.000 - R$ 1.800.000 (±10-15% incerteza)",
          size=10)
    doc.add_paragraph()

    add_h(doc, "9.3. Coerência com o orçamento automatizado", size=12)
    add_p(doc,
          "Os valores acima estão CONSOLIDADOS dentro dos macrogrupos correspondentes "
          "do executivo automatizado:")
    add_p(doc,
          "• AC e exaustão entram no macrogrupo Climatização (R$ 51,02/m² × 37.894 m² = R$ 1.933.000) "
          "e parcialmente em Sistemas Especiais", size=10)
    add_p(doc,
          "• A estimativa BIM da exaustão (R$ 1,1-1,8M) e do AC (138 TR × R$ 6.000-9.000/TR ≈ "
          "R$ 828k-1,2M) total em R$ 2-3M, levemente acima do calibrado (R$ 1,9M de Climatização)", size=10)
    add_p(doc,
          "• Sugere-se ajuste fino manual no executivo para refletir os valores BIM exatos", size=10)
    doc.add_paragraph()

    add_h(doc, "9.4. Próximas extrações pendentes", size=12)
    add_p(doc, "Conforme RESUMO-EXTRACAO.md do BIM, faltam ainda extrair:", size=10)
    add_p(doc, "• Tubulações frigoríficas (layers específicos)", size=10)
    add_p(doc, "• Linhas de dreno", size=10)
    add_p(doc, "• Instalações elétricas dos splits", size=10)
    add_p(doc, "• Suportes e acessórios (atualmente estimados por quantidade de equipamentos)", size=10)

    doc.save(str(docx_path))
    print(f"  OK: {docx_path.name} patchado")
    return True


if __name__ == "__main__":
    print("=== Patching Thozen memoriais com dados BIM ===")
    print(f"\n[1/2] Memorial executivo:")
    patch(EXEC_DOCX)
    print(f"\n[2/2] Memorial paramétrico:")
    patch(PARAM_DOCX)
    print("\nDone.")
