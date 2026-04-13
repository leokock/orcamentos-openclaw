#!/usr/bin/env python3
"""Gera memorial Word (.docx) consolidado do pacote (paramétrico OU executivo).

Diferente do gerar_memorial_rastreavel.py original (que parseia uma planilha
de template específico), este script monta o memorial a partir dos dados
consolidados no `state.json` do pacote + camada qualitativa Gemma + análise
arquitetônica do Bloco 0.

Uso:
    python gerar_memorial_pacote.py --slug projeto --tipo parametrico
    python gerar_memorial_pacote.py --slug projeto --tipo executivo
"""
from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).parent))
import consulta_similares as cs  # noqa: E402

BASE = Path.home() / "orcamentos-openclaw" / "base"
PACOTES = BASE / "pacotes"


def fmt_money(v) -> str:
    if v is None:
        return "—"
    try:
        return f"R$ {float(v):,.0f}".replace(",", ".")
    except Exception:
        return str(v)


def fmt_num(v, dec=2) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v):,.{dec}f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return str(v)


def style_heading(p, size=16, bold=True, color="2C3E50"):
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    sizes = {0: 22, 1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(sizes.get(level, 11))
    run.font.bold = True
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string("2C3E50")
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_para(doc, text, size=10, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_table(doc, headers, rows, widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor.from_string("FFFFFF")

    for row in rows:
        r = table.add_row()
        for i, v in enumerate(row[:n_cols]):
            c = r.cells[i]
            c.text = str(v) if v is not None else "—"
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"
    return table


def load_state(slug: str) -> dict:
    p = PACOTES / slug / "state.json"
    if not p.exists():
        raise FileNotFoundError(f"state.json não encontrado em {p}")
    return json.loads(p.read_text(encoding="utf-8"))


def load_arquitetura(slug: str) -> dict:
    p = PACOTES / slug / "analise-arquitetura.json"
    if not p.exists():
        return {}
    return json.loads(p.read_text(encoding="utf-8"))


def gate_decisoes(slug: str) -> dict:
    """Lê o gate validado se existir."""
    pasta = PACOTES / slug
    candidates = list(pasta.glob(f"gate-{slug}-validado.xlsx"))
    if not candidates:
        return {}
    try:
        from openpyxl import load_workbook
        wb = load_workbook(candidates[0], data_only=True)
        if "GATE" not in wb.sheetnames:
            return {}
        ws = wb["GATE"]
        out = {}
        for row in ws.iter_rows(min_row=5, values_only=True):
            if row and row[0] and row[1]:
                out[str(row[0]).strip()] = str(row[1]).strip()
        wb.close()
        return out
    except Exception:
        return {}


def gerar(slug: str, tipo: str, output: Path) -> dict:
    state = load_state(slug)
    arquit = load_arquitetura(slug)
    decisoes = gate_decisoes(slug)

    ac = state["ac"]
    ur = state.get("ur")
    padrao = state.get("padrao")

    similares = cs.projetos_similares(ac=ac, ur=ur, padrao=padrao, n=5)
    valores_calib = cs.valores_macrogrupos_calibrados(ac, padrao=padrao)
    grand_total = sum(d.get("total_estimado", 0) for d in valores_calib.values())

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    title = "MEMORIAL DESCRITIVO" if tipo == "parametrico" else "MEMORIAL EXECUTIVO"
    sub = "Orçamento Paramétrico V2" if tipo == "parametrico" else "Orçamento Executivo Automatizado"
    add_heading(doc, title, level=0)
    add_heading(doc, sub, level=1)
    add_para(doc, f"Projeto: {slug.upper()}", size=12, bold=True)
    add_para(doc, f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}", italic=True)
    doc.add_paragraph()

    add_heading(doc, "1. Dados do Projeto", level=1)
    rows = [
        ("Slug", slug),
        ("AC (área construída)", f"{fmt_num(ac, 2)} m²"),
        ("UR (unidades residenciais)", str(ur or "—")),
        ("Padrão de acabamento", padrao or "—"),
    ]
    add_table(doc, ["Campo", "Valor"], rows)
    doc.add_paragraph()

    add_heading(doc, "2. Premissas Validadas", level=1)
    if decisoes:
        add_para(doc, f"Decisões aprovadas no gate de validação ({len(decisoes)} itens):", italic=True)
        rows = [(k, v) for k, v in decisoes.items()]
        add_table(doc, ["Decisão", "Valor adotado"], rows)
    else:
        add_para(doc, "(gate validado não encontrado — usando defaults)", italic=True)
    doc.add_paragraph()

    if arquit and arquit.get("decisoes_inferidas"):
        add_heading(doc, "3. Análise Arquitetônica (Bloco 0)", level=1)
        add_para(doc,
                 f"Análise multi-camada (IFC + DXF + PDF + nomes) executada em {arquit.get('duracao_s', '?')}s.",
                 italic=True)
        decisoes_arq = arquit.get("decisoes_inferidas", {})
        rows = []
        labels = {
            "tem_piscina": "Piscina",
            "piscina_aquecida": "Piscina aquecida",
            "tem_ofuro_spa": "Ofurô / SPA",
            "tem_sauna": "Sauna",
            "tem_academia": "Academia",
            "tem_quadra": "Quadra esportiva",
            "tem_salao_festas": "Salão de festas",
            "tem_gourmet": "Espaço gourmet",
            "tem_churrasqueira": "Churrasqueira",
            "tem_playground": "Playground / brinquedoteca",
            "tem_coworking": "Coworking",
            "tem_pet": "Pet place",
            "tem_bicicletario": "Bicicletário",
            "tem_gerador_dedicado": "Gerador dedicado",
        }
        for key, label in labels.items():
            rows.append((label, "✓ Sim" if decisoes_arq.get(key) else "—"))
        add_table(doc, ["Item de lazer/sistema", "Detectado no projeto"], rows)
        doc.add_paragraph()

    add_heading(doc, f"{4 if arquit and arquit.get('decisoes_inferidas') else 3}. Composição de Custos por Macrogrupo", level=1)
    add_para(doc,
             f"Custos calibrados a partir de {sum(d.get('n_amostras', 0) for d in valores_calib.values())} "
             f"amostras de {len(set().union(*[set(d.get('fontes', [])) for d in valores_calib.values()]))} "
             f"projetos similares.",
             italic=True)

    rows = []
    for mg, dados in valores_calib.items():
        total = dados.get("total_estimado", 0)
        rsm2 = dados.get("rsm2_ajustado", dados.get("rsm2_mediano", 0))
        pct = total / grand_total * 100 if grand_total else 0
        n = dados.get("n_amostras", 0)
        rows.append((mg, fmt_money(total), fmt_num(rsm2, 2), f"{pct:.1f}%", str(n)))
    rows.append(("TOTAL", fmt_money(grand_total), fmt_num(grand_total / ac, 2) if ac else "—", "100%", "—"))
    add_table(doc, ["Macrogrupo", "Total R$", "R$/m²", "%", "N amostras"], rows)
    doc.add_paragraph()

    add_heading(doc, "5. Projetos Similares Utilizados", level=1)
    add_para(doc, f"{len(similares)} projetos da base Cartesian usados como referência:", italic=True)
    rows = []
    for p in similares:
        rows.append((
            p["_slug"],
            f"{fmt_num(p.get('ac'), 0)} m²",
            str(p.get("ur") or "—"),
            fmt_money(p.get("total")),
            fmt_num(p.get("rsm2"), 2) if p.get("rsm2") else "—",
        ))
    add_table(doc, ["Projeto", "AC", "UR", "Total R$", "R$/m²"], rows)
    doc.add_paragraph()

    add_heading(doc, "6. Sub-disciplinas (Camada Gemma)", level=1)
    enriq = cs.enriquecer_parametrico(similares)
    sub_por_mg = enriq.get("sub_disciplinas_por_mg", {})
    if sub_por_mg:
        add_para(doc,
                 f"Sub-disciplinas reais identificadas via análise qualitativa Gemma "
                 f"em {enriq.get('n_similares', 0)} projetos similares:",
                 italic=True)
        for mg, subs in sub_por_mg.items():
            if not subs:
                continue
            add_heading(doc, mg, level=3)
            for sd in subs[:6]:
                add_para(doc,
                         f"• {sd['sub']} (freq {sd['freq']}/{enriq['n_similares']}, "
                         f"fontes: {', '.join(sd['fontes'][:3])})",
                         size=9)
    else:
        add_para(doc, "(sem sub-disciplinas extraídas para esses similares)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "7. Premissas Técnicas Consolidadas", level=1)
    premissas = cs.premissas_consolidadas(similares)
    if premissas:
        add_para(doc, f"{len(premissas)} premissas extraídas dos memoriais dos similares:", italic=True)
        rows = []
        for pr in premissas[:25]:
            rows.append((
                pr.get("area", ""),
                pr.get("premissa", "")[:80],
                f"{pr.get('freq', 1)}/{enriq['n_similares']}",
            ))
        add_table(doc, ["Área", "Premissa", "Frequência"], rows)
    else:
        add_para(doc, "(nenhuma premissa consolidada disponível)", italic=True)
    doc.add_paragraph()

    add_heading(doc, "8. Validação", level=1)
    seg_label = (
        "Pequeno (<8k m²)" if ac < 8000 else
        "Médio (8-15k m²)" if ac < 15000 else
        "Grande (15-25k m²)" if ac < 25000 else
        "Extra (>25k m²)"
    )
    add_para(doc, f"Segmento por porte: {seg_label}", bold=True)
    add_para(doc, f"R$/m² estimado: {fmt_num(grand_total / ac if ac else 0, 2)}")
    add_para(doc, f"Total estimado: {fmt_money(grand_total)}")
    doc.add_paragraph()

    add_para(doc,
             f"Memorial gerado automaticamente por gerar_memorial_pacote.py em "
             f"{datetime.now().strftime('%d/%m/%Y %H:%M')}",
             size=8, italic=True)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

    return {
        "output": str(output),
        "n_macrogrupos": len(valores_calib),
        "grand_total": round(grand_total, 2),
        "rsm2": round(grand_total / ac, 2) if ac else 0,
        "n_similares": len(similares),
        "n_premissas": len(premissas),
        "n_sub_disciplinas": sum(len(s) for s in sub_por_mg.values()),
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--slug", required=True)
    ap.add_argument("--tipo", choices=["parametrico", "executivo"], default="parametrico")
    ap.add_argument("-o", "--output", default=None)
    args = ap.parse_args()

    if args.output:
        out = Path(args.output)
    else:
        out = PACOTES / args.slug / f"{args.tipo}-{args.slug}.docx"

    result = gerar(args.slug, args.tipo, out)
    print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
