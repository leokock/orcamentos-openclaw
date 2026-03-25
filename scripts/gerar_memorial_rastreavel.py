#!/usr/bin/env python3
"""
Gera memorial Word com rastreabilidade per-item a partir da planilha grande.

Modelo C+A: le a planilha finalizada e gera o docx automaticamente.
Cada item mostra de onde veio (projetista, versao do projeto, base parametrica, etc).

Uso:
  python gerar_memorial_rastreavel.py <planilha.xlsx> [--projetistas projetistas.json] [--output memorial.docx]

Se --projetistas nao for fornecido, usa tags genericas (IFC/DXF, Referencia, Estimado).
"""
import argparse
import json
import os
import sys
import unicodedata
from collections import Counter

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import openpyxl

# === CORES ===
GREEN_HEX = "C6EFCE"
YELLOW_HEX = "FFEB9C"
RED_HEX = "FFC7CE"
GRAY_HEX = "F0F0F0"
HEADER_HEX = "2C3E50"

DARK_BLUE = RGBColor(0x2C, 0x3E, 0x50)
ACCENT = RGBColor(0x29, 0x80, 0xB9)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

COLOR_MAP = {
    '00C6EFCE': ('IFC/DXF', GREEN_HEX),
    '00FFEB9C': ('Referencia', YELLOW_HEX),
    '00FFC7CE': ('Estimado', RED_HEX),
}


# === HELPERS ===
def set_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def brl(v):
    if v is None:
        return "-"
    return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def brl_int(v):
    if v is None:
        return "-"
    return f"R$ {v:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")


def get_cell_source(cell):
    fill = cell.fill
    for attr in ['fgColor', 'start_color']:
        c = getattr(fill, attr, None)
        if c and c.rgb and c.rgb in COLOR_MAP:
            return COLOR_MAP[c.rgb]
    return ('Sem tag', GRAY_HEX)


def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE if level == 1 else ACCENT
    return h


def add_header_row(table, cols):
    row = table.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        run.font.name = "Arial"
        set_shading(cell, HEADER_HEX)


def sanitize_origin(text):
    """Remove nomes de projetos de outros clientes."""
    # Adicionar nomes de projetos que devem ser sanitizados
    CONFIDENTIAL = [
        'Elizabeth II', 'elizabeth II', 'Elizabeth 2',
        'Maison Beach', 'Catena', 'Connect', 'SOHO',
    ]
    result = text
    for name in CONFIDENTIAL:
        result = result.replace(name, 'projeto referencia similar')
    return result


# === EXTRACT DATA ===
def extract_capa(wb):
    """Extrai dados do empreendimento da aba CAPA."""
    ws = wb['CAPA']
    data = {}
    field_map = {
        'Projeto': 'projeto', 'Empresa': 'empresa', 'Revisão': 'revisao',
        'Prazo de Obra': 'prazo', 'Área do Terreno': 'area_terreno',
        'Área Construída': 'ac', 'Área Privativa': 'area_privativa',
        'Área de Lazer': 'area_lazer', 'Área Comercial': 'area_comercial',
    }
    field_map_e = {
        'Nº Unid. Residenciais': 'ur', 'Nº Unid. Comerciais': 'uc',
        'Nº Pavimentos': 'pavimentos', 'Nº Subsolos': 'subsolos',
        'Nº Vagas': 'vagas', 'ENDEREÇO DA OBRA': 'endereco',
        'Cidade': 'cidade',
    }
    for r in range(1, min(25, ws.max_row + 1)):
        for c in range(1, min(10, ws.max_column + 1)):
            v = ws.cell(r, c).value
            if not v or not isinstance(v, str):
                continue
            v_clean = v.strip().rstrip(':')
            for key, field in {**field_map, **field_map_e}.items():
                if key in v_clean:
                    for cc in range(c + 1, min(c + 3, ws.max_column + 1)):
                        val = ws.cell(r, cc).value
                        if val is not None:
                            data[field] = val
                            break
    return data


def extract_eap(wb):
    """Extrai EAP da aba 'EAP Analise Revisao'."""
    sheet_name = None
    for sn in wb.sheetnames:
        if 'EAP' in sn and ('Analise' in sn or 'Análise' in sn):
            sheet_name = sn
            break
    if not sheet_name:
        return []

    ws = wb[sheet_name]
    rows = []
    for r in range(3, ws.max_row + 1):
        etapa = ws.cell(r, 1).value
        valor = ws.cell(r, 2).value
        pct = ws.cell(r, 3).value
        rsm2 = ws.cell(r, 4).value
        obs = ws.cell(r, 9).value
        if not etapa or not valor or not isinstance(valor, (int, float)):
            continue
        if 'TOTAL' in str(etapa).upper() or 'LEGENDA' in str(etapa).upper():
            continue
        rows.append({
            'etapa': str(etapa),
            'valor': float(valor),
            'pct': float(pct) if pct and isinstance(pct, (int, float)) else 0,
            'rsm2': float(rsm2) if rsm2 and isinstance(rsm2, (int, float)) else 0,
            'obs': sanitize_origin(str(obs)) if obs else '',
        })
    return rows


def extract_discipline_items(wb, projetistas):
    """Extrai itens de todas as abas de disciplina."""
    # Detect discipline tabs: have Origem row and standard column structure
    SKIP_TABS = {'CAPA', 'EAP Thozen', 'EAP Análise Revisão', 'EAP Analise Revisao',
                 'Ger_Executivo_Cartesian', 'Cálculo de apoio', 'Escoramento',
                 'M.O. ESTRUTURA', 'MÃO DE OBRA ELETRICA', 'MÃO DE OBRA HIDRO',
                 'PROPOSTA CPUs', 'Resumo Estrutura'}

    all_tabs = {}
    for sheet_name in wb.sheetnames:
        if sheet_name in SKIP_TABS:
            continue
        ws = wb[sheet_name]
        if ws.max_row < 7:
            continue

        # Check if it has the standard discipline structure (col E=desc, L=total)
        has_origem = False
        origem = ""
        for r in range(1, 5):
            for c in range(1, 10):
                v = ws.cell(r, c).value
                if v and isinstance(v, str) and 'Origem' in v:
                    has_origem = True
                    for cc in range(c + 1, c + 3):
                        ov = ws.cell(r, cc).value
                        if ov:
                            origem = sanitize_origin(str(ov))
                            break

        if not has_origem:
            continue

        # Get projetista for this discipline
        projetista = projetistas.get(sheet_name, '')

        # Extract items
        items = []
        for r in range(7, ws.max_row + 1):
            desc = ws.cell(r, 5).value
            total = ws.cell(r, 12).value
            if not desc or not total or not isinstance(total, (int, float)):
                continue
            desc_str = str(desc)
            if any(x in desc_str.upper() for x in ['R$/M', 'LEGENDA', 'TOTAL GERAL']):
                continue

            qty = ws.cell(r, 9).value or ws.cell(r, 6).value
            unit = ws.cell(r, 10).value
            pu = ws.cell(r, 11).value
            generic_tag, color = get_cell_source(ws.cell(r, 5))

            # Build specific tag
            if color == GREEN_HEX:
                proj_prefix = f"Proj. {projetista} " if projetista else "Proj. "
                if '— DXF' in desc_str:
                    for part in origem.split('+'):
                        part = part.strip()
                        if 'DXF' in part:
                            tag = f"{proj_prefix}{part}".strip()[:28]
                            break
                    else:
                        tag = f"{proj_prefix}DXF".strip()[:28]
                elif '— BIM' in desc_str:
                    tag = f"{proj_prefix}IFC (BIM)".strip()[:28]
                else:
                    tag = proj_prefix + "IFC"
                    for part in origem.split('+'):
                        part = part.strip()
                        if 'IFC' in part and ('rev' in part.lower() or 'R' in part):
                            tag = f"{proj_prefix}{part}".strip()[:28]
                            break
                source_tag = tag
            elif color == YELLOW_HEX:
                dl = desc_str.lower()
                if 'sinapi' in dl:
                    source_tag = 'SINAPI'
                elif 'base' in dl and 'R$' in desc_str:
                    source_tag = 'PU base Cartesian'
                elif 'estimado' in dl or 'indice' in dl:
                    source_tag = 'Indice param.'
                elif projetista and 'briefing' in dl:
                    source_tag = f'Briefing {projetista}'
                else:
                    source_tag = 'Param. base Cartesian'
            elif color == RED_HEX:
                source_tag = 'Estimado (s/ projeto)'
            else:
                source_tag = generic_tag

            items.append({
                'desc': desc_str[:70],
                'qty': float(qty) if qty and isinstance(qty, (int, float)) else None,
                'unit': str(unit) if unit else '',
                'pu': float(pu) if pu and isinstance(pu, (int, float)) else None,
                'total': float(total),
                'source_tag': source_tag,
                'source_color': color,
            })

        if items:
            tab_total = sum(it['total'] for it in items)
            all_tabs[sheet_name] = {
                'origem': origem,
                'projetista': projetista,
                'items': items,
                'total': tab_total,
            }
            print(f"  {sheet_name}: {len(items)} itens | R$ {tab_total:,.0f} | Proj: {projetista or '-'}")

    return all_tabs


def extract_ger_items(wb):
    """Extrai itens do Gerenciamento (formato diferente)."""
    if 'Ger_Tec e Adm' not in wb.sheetnames:
        return []

    ws = wb['Ger_Tec e Adm']
    items = []
    for r in range(9, ws.max_row + 1):
        nivel = ws.cell(r, 8).value
        desc = ws.cell(r, 10).value
        unit = ws.cell(r, 11).value
        qty = ws.cell(r, 12).value
        pu = ws.cell(r, 13).value
        total = ws.cell(r, 14).value
        obs = str(ws.cell(r, 15).value or '')

        if not desc or not total or not isinstance(total, (int, float)):
            continue
        if nivel in ['CELULA CONSTRUTIVA', 'ETAPA', 'CÉLULA CONSTRUTIVA']:
            continue

        obs_lower = obs.lower()
        if 'contrato' in obs_lower:
            tag = 'Contrato'
            for name in ['Battisti', 'Favola', 'DMA', 'Zeplin', 'Liberte',
                         'franzmann', 'value', 'W.Thomaz', 'WM', 'otus', 'triunfo']:
                if name.lower() in obs_lower:
                    tag = f'Contrato {name.title()}'
                    break
            source_tag, source_color = tag, GREEN_HEX
        elif 'repessado' in obs_lower or 'gessele' in obs_lower:
            # Use project name from CAPA, not client's internal name
            source_tag, source_color = 'Confirmado cliente', GREEN_HEX
        elif 'estimado' in obs_lower:
            source_tag, source_color = 'Estimado (s/ proj)', RED_HEX
        elif 'sinapi' in obs_lower:
            source_tag, source_color = 'SINAPI', YELLOW_HEX
        else:
            source_tag, source_color = 'Param. base Cartesian', YELLOW_HEX

        items.append({
            'desc': str(desc)[:70],
            'qty': float(qty) if qty and isinstance(qty, (int, float)) else None,
            'unit': str(unit) if unit else '',
            'pu': float(pu) if pu and isinstance(pu, (int, float)) else None,
            'total': float(total),
            'source_tag': source_tag,
            'source_color': source_color,
        })

    print(f"  Ger_Tec e Adm: {len(items)} itens")
    return items


def extract_estrutura(wb):
    """Extrai resumo estrutural."""
    if 'Resumo Estrutura' not in wb.sheetnames:
        return []
    ws = wb['Resumo Estrutura']
    items = []
    for r in range(4, ws.max_row + 1):
        disc = ws.cell(r, 2).value
        if not disc:
            continue
        items.append({
            'disc': str(disc),
            'pav': str(ws.cell(r, 3).value or ''),
            'etapa': str(ws.cell(r, 4).value or ''),
            'proj': str(ws.cell(r, 5).value or ''),
            'concreto': str(ws.cell(r, 8).value or ''),
            'forma': str(ws.cell(r, 9).value or ''),
        })
    return items


# === GENERATE DOCUMENT ===
def generate_doc(capa, eap_rows, all_tabs, ger_items, est_items, output_path):
    ac = float(capa.get('ac', 1))
    projeto = capa.get('projeto', 'Projeto')
    empresa = capa.get('empresa', 'Empresa')
    total_geral = sum(e['valor'] for e in eap_rows)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(3)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # === CAPA ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CARTESIAN ENGENHARIA")
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ORCAMENTO EXECUTIVO")
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(projeto).upper())
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Revisao {capa.get('revisao', '')} | Rastreabilidade Completa")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    doc.add_paragraph()

    # Dados do projeto
    dados_pairs = [
        ("Empresa", str(empresa)),
        ("Cidade/UF", str(capa.get('cidade', capa.get('endereco', '')))),
        ("AC", f"{ac:,.2f} m2".replace(",", "X").replace(".", ",").replace("X", ".")),
        ("Pavimentos", str(capa.get('pavimentos', ''))),
        ("UR", str(capa.get('ur', ''))),
        ("Vagas", str(capa.get('vagas', ''))),
        ("Prazo", f"{capa.get('prazo', '')} meses"),
    ]
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    for i in range(0, len(dados_pairs), 2):
        row = table.add_row()
        for j in range(2):
            if i + j < len(dados_pairs):
                k, v = dados_pairs[i + j]
                row.cells[j * 2].text = k
                row.cells[j * 2 + 1].text = v
                for c in [row.cells[j * 2], row.cells[j * 2 + 1]]:
                    for run in c.paragraphs[0].runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Arial'

    doc.add_page_break()

    # === 1. LEGENDA ===
    add_styled_heading(doc, "1. Legenda de Rastreabilidade", level=1)

    table_leg = doc.add_table(rows=4, cols=3)
    table_leg.style = 'Table Grid'
    add_header_row(table_leg, ["Cor", "Classificacao", "Significado"])
    leg_data = [
        (GREEN_HEX, "Projeto / Confirmado", f"Extraido de IFC/DXF dos projetistas do {projeto} ou confirmado pelo cliente"),
        (YELLOW_HEX, "Parametrico / Referencia", "PU ou qtd baseado na base Cartesian (75 executivos) ou SINAPI"),
        (RED_HEX, "Estimado (sem projeto)", "Sem fonte confiavel - requer cotacao ou projeto para proxima revisao"),
    ]
    for i, (color, tag, desc) in enumerate(leg_data):
        row = table_leg.rows[i + 1]
        set_shading(row.cells[0], color)
        row.cells[0].text = ""
        row.cells[1].text = tag
        row.cells[2].text = desc
        for c in row.cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(8.5)
                run.font.name = 'Arial'

    doc.add_paragraph()

    # === 2. EAP RESUMO ===
    add_styled_heading(doc, "2. EAP Resumo com Origem", level=1)

    table_eap = doc.add_table(rows=1, cols=5)
    table_eap.style = 'Table Grid'
    add_header_row(table_eap, ["Etapa", "Valor (R$)", "R$/m2", "%", "Origem"])

    for e in eap_rows:
        row = table_eap.add_row()
        row.cells[0].text = e['etapa'][:45]
        row.cells[1].text = brl_int(e['valor'])
        row.cells[2].text = f"{e['rsm2']:.0f}"
        row.cells[3].text = f"{e['pct'] * 100:.1f}%"
        row.cells[4].text = e['obs'][:40]
        for c in row.cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(7.5)
                run.font.name = 'Arial'
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # === 3. ESTRUTURA ===
    if est_items:
        add_styled_heading(doc, "3. Estrutura - Quantitativos", level=1)
        table_est = doc.add_table(rows=1, cols=6)
        table_est.style = 'Table Grid'
        add_header_row(table_est, ["Disciplina", "Pavimento", "Etapa", "Projetos", "Concreto (m3)", "Forma (m2)"])
        for it in est_items:
            row = table_est.add_row()
            for i, val in enumerate([it['disc'], it['pav'], it['etapa'], it['proj'], it['concreto'], it['forma']]):
                row.cells[i].text = val
                set_shading(row.cells[i], GREEN_HEX)
                for run in row.cells[i].paragraphs[0].runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Arial'
        doc.add_paragraph()

    # === 4. GERENCIAMENTO ===
    if ger_items:
        section_num = 4
        add_styled_heading(doc, f"{section_num}. Gerenciamento Tecnico e Administrativo", level=1)
        p = doc.add_paragraph()
        p.add_run("Valores confirmados pelo cliente (verde), parametricos (amarelo) ou estimados (vermelho).").font.color.rgb = GRAY
        _add_items_table(doc, ger_items)
        section_num = 5
    else:
        section_num = 4

    # === DISCIPLINAS ===
    for sheet_name, tab_data in all_tabs.items():
        items = tab_data['items']
        if not items:
            continue

        add_styled_heading(doc, f"{section_num}. {sheet_name}", level=1)

        p = doc.add_paragraph()
        run = p.add_run("Origem: ")
        run.font.bold = True
        p.add_run(tab_data['origem'] or "Ver legenda")
        if tab_data['projetista']:
            p = doc.add_paragraph()
            run = p.add_run("Projetista: ")
            run.font.bold = True
            p.add_run(tab_data['projetista'])

        p = doc.add_paragraph()
        p.add_run(f"Total: {brl_int(tab_data['total'])} | {len(items)} itens | R$/m2: {tab_data['total'] / ac:.2f}")

        # Source distribution
        src_count = Counter(it['source_tag'] for it in items)
        src_val = {}
        for it in items:
            src_val[it['source_tag']] = src_val.get(it['source_tag'], 0) + it['total']
        p = doc.add_paragraph()
        parts = []
        for tag, count in src_count.most_common():
            val = src_val[tag]
            parts.append(f"{tag}: {count} itens ({val / tab_data['total'] * 100:.0f}%)")
        p.add_run("  |  ".join(parts)).font.size = Pt(8.5)

        _add_items_table(doc, items)
        section_num += 1

    # === PREMISSAS ===
    doc.add_page_break()
    add_styled_heading(doc, f"{section_num}. Premissas e Limitacoes", level=1)

    premissas = [
        f"Area construida: {ac:,.2f} m2 (projeto arquitetonico)",
        "Parametricos: mediana base Cartesian (75 executivos calibrados)",
        "Valores nao incluem terreno, incorporacao, marketing ou despesas financeiras",
    ]
    for item in premissas:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Cartesian Engenharia")
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    doc.save(output_path)
    print(f"\nGerado: {output_path}")


def _add_items_table(doc, items):
    """Adiciona tabela de itens com rastreabilidade."""
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    add_header_row(table, ["Descricao", "Qtd", "Unid", "PU (R$)", "Total (R$)", "Fonte"])

    for it in items:
        row = table.add_row()
        qty_str = f"{it['qty']:,.1f}" if it['qty'] else "-"
        pu_str = brl(it['pu']) if it['pu'] else "-"

        cells = [
            (it['desc'], WD_ALIGN_PARAGRAPH.LEFT),
            (qty_str, WD_ALIGN_PARAGRAPH.CENTER),
            (it['unit'], WD_ALIGN_PARAGRAPH.CENTER),
            (pu_str, WD_ALIGN_PARAGRAPH.RIGHT),
            (brl_int(it['total']), WD_ALIGN_PARAGRAPH.RIGHT),
        ]
        for i, (text, align) in enumerate(cells):
            row.cells[i].text = str(text)
            row.cells[i].paragraphs[0].alignment = align
            for run in row.cells[i].paragraphs[0].runs:
                run.font.size = Pt(7.5)
                run.font.name = 'Arial'

        row.cells[5].text = it['source_tag']
        row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_shading(row.cells[5], it['source_color'])
        for run in row.cells[5].paragraphs[0].runs:
            run.font.size = Pt(7)
            run.font.name = 'Arial'
            run.font.bold = True


# === MAIN ===
def main():
    parser = argparse.ArgumentParser(description='Gera memorial Word rastreavel a partir da planilha grande.')
    parser.add_argument('xlsx', help='Planilha grande (.xlsx)')
    parser.add_argument('--projetistas', help='JSON com mapeamento de projetistas por disciplina')
    parser.add_argument('--output', '-o', help='Arquivo de saida (.docx)')
    args = parser.parse_args()

    xlsx_path = os.path.expanduser(args.xlsx)
    if not os.path.exists(xlsx_path):
        print(f"Erro: {xlsx_path} nao encontrado")
        sys.exit(1)

    # Load projetistas mapping
    projetistas = {}
    if args.projetistas:
        with open(os.path.expanduser(args.projetistas)) as f:
            projetistas = json.load(f)
        print(f"Projetistas carregados: {len(projetistas)} disciplinas")

    # Output path
    if args.output:
        output_path = os.path.expanduser(args.output)
    else:
        base = os.path.splitext(xlsx_path)[0]
        output_path = f"{base}-MEMORIAL-RASTREAVEL.docx"

    print(f"Lendo {xlsx_path}...")
    wb = openpyxl.load_workbook(xlsx_path)

    capa = extract_capa(wb)
    print(f"Projeto: {capa.get('projeto', '?')} | AC: {capa.get('ac', '?')} m2")

    eap_rows = extract_eap(wb)
    print(f"EAP: {len(eap_rows)} linhas")

    est_items = extract_estrutura(wb)
    ger_items = extract_ger_items(wb)
    all_tabs = extract_discipline_items(wb, projetistas)

    total_items = sum(len(t['items']) for t in all_tabs.values()) + len(ger_items) + len(est_items)
    print(f"\nTotal de itens rastreados: {total_items}")

    generate_doc(capa, eap_rows, all_tabs, ger_items, est_items, output_path)


if __name__ == '__main__':
    main()
