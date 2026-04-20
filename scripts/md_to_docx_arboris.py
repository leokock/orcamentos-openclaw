"""
Converte os .md de Premissas-Origem e Justificativa-Itens-Acima-Media em .docx
usando python-docx com estilo Cartesian.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL_PRIM = RGBColor(0x1E, 0x52, 0xF5)
AZUL_NAVY = RGBColor(0x0F, 0x1C, 0x4A)
LARANJA = RGBColor(0xFF, 0x57, 0x22)
CINZA_TXT = RGBColor(0x45, 0x4E, 0x5E)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = AZUL_PRIM


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = AZUL_NAVY


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = AZUL_PRIM


def parse_inline(text):
    """Retorna [(chunk, bold, italic, code)] — remove emojis antes."""
    text = _strip_emoji(text)
    out = []
    i = 0
    current = ''
    while i < len(text):
        if text[i:i+2] == '**':
            if current:
                out.append((current, False, False, False))
                current = ''
            j = text.find('**', i + 2)
            if j < 0:
                current += text[i:]
                i = len(text)
            else:
                out.append((text[i+2:j], True, False, False))
                i = j + 2
        elif text[i] == '`':
            if current:
                out.append((current, False, False, False))
                current = ''
            j = text.find('`', i + 1)
            if j < 0:
                current += text[i:]
                i = len(text)
            else:
                out.append((text[i+1:j], False, False, True))
                i = j + 1
        elif text[i:i+2] == '~~':
            if current:
                out.append((current, False, False, False))
                current = ''
            j = text.find('~~', i + 2)
            if j < 0:
                current += text[i:]
                i = len(text)
            else:
                out.append((text[i+2:j], False, True, False))
                i = j + 2
        else:
            current += text[i]
            i += 1
    if current:
        out.append((current, False, False, False))
    return out


def add_para(doc, text, size=11, bold=False, italic=False, color=None, quote=False):
    p = doc.add_paragraph()
    if quote:
        p.paragraph_format.left_indent = Inches(0.3)
    runs = parse_inline(text)
    for chunk, is_bold, is_italic, is_code in runs:
        r = p.add_run(chunk)
        r.font.name = 'Consolas' if is_code else 'Calibri'
        r.font.size = Pt(size)
        r.font.bold = bold or is_bold
        r.font.italic = italic or is_italic
        if color:
            r.font.color.rgb = color
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    runs = parse_inline(text)
    for chunk, is_bold, is_italic, is_code in runs:
        r = p.add_run(chunk)
        r.font.name = 'Consolas' if is_code else 'Calibri'
        r.font.size = Pt(size)
        r.font.bold = is_bold
        r.font.italic = is_italic


EMOJI_RE = re.compile(
    '[' +
    '\U0001F300-\U0001FAFF' +  # symbols & pictographs, supplemental
    '\U0001F600-\U0001F64F' +  # emoticons
    '\U0001F680-\U0001F6FF' +  # transport
    '\U0001F700-\U0001F77F' +
    '\U0001F780-\U0001F7FF' +
    '\U0001F800-\U0001F8FF' +
    '\U0001F900-\U0001F9FF' +
    '\U0001FA00-\U0001FA6F' +
    '\U0001FA70-\U0001FAFF' +
    '\u2600-\u26FF' +          # misc symbols
    '\u2700-\u27BF' +          # dingbats
    '\u2190-\u21FF' +          # arrows
    '\u2B00-\u2BFF' +
    '\uFE0F' +                 # variation selector
    ']',
    flags=re.UNICODE
)


def _strip_emoji(s):
    """Remove emojis/marcadores e normaliza espaços."""
    return re.sub(r'\s+', ' ', EMOJI_RE.sub('', s)).strip()


def add_table_md(doc, rows):
    if not rows:
        return
    header = rows[0]
    data_rows = []
    for r in rows[1:]:
        if all(set(c.strip()) <= set('-: ') for c in r):
            continue
        data_rows.append(r)
    n_cols = len(header)

    # Se a última coluna fica vazia após remover emojis, descarta ela
    if n_cols >= 2:
        header_last_clean = _strip_emoji(header[-1])
        data_last_clean = [_strip_emoji(r[-1]) if len(r) >= n_cols else '' for r in data_rows]
        # Se coluna vazia em quase todas as linhas (após strip emoji): descartar
        if all(not c or len(c) <= 2 for c in data_last_clean):
            header = header[:-1]
            data_rows = [r[:-1] for r in data_rows]
            n_cols -= 1

    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.autofit = False

    # Larguras: primeira coluna menor, segunda maior (pra caber texto)
    if n_cols == 2:
        widths = [Cm(6.5), Cm(10.5)]
    elif n_cols == 3:
        widths = [Cm(5.0), Cm(7.5), Cm(4.5)]
    elif n_cols == 4:
        widths = [Cm(4.5), Cm(4.5), Cm(4.0), Cm(4.0)]
    else:
        widths = [Cm(17.0 / n_cols)] * n_cols

    for j, w in enumerate(widths):
        for row in table.rows:
            row.cells[j].width = w

    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        for chunk, b, it, code in parse_inline(h.strip()):
            r = p.add_run(chunk)
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, '1E52F5')

    for i, row in enumerate(data_rows):
        for j, val in enumerate(row[:n_cols]):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            stripped = val.strip()
            if stripped and (stripped.startswith('R$') or
                             (stripped.replace('.', '').replace(',', '').replace('%', '').replace('+', '').replace('-', '').replace(' ', '').replace('*', '').isdigit())):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for chunk, b, it, code in parse_inline(stripped):
                r = p.add_run(chunk)
                r.font.name = 'Consolas' if code else 'Calibri'
                r.font.size = Pt(9)
                r.font.bold = b
                r.font.italic = it
            if i % 2 == 1:
                set_cell_bg(cell, 'F0F4F9')


def add_yaml_block(doc, lines):
    add_h3(doc, 'Metadados')
    for ln in lines:
        if ':' in ln:
            k, _, v = ln.partition(':')
            p = doc.add_paragraph()
            r = p.add_run(f'{k.strip()}: ')
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = AZUL_NAVY
            for chunk, b, it, code in parse_inline(v.strip()):
                r2 = p.add_run(chunk)
                r2.font.name = 'Calibri'
                r2.font.size = Pt(10)
                r2.font.color.rgb = CINZA_TXT
                r2.font.bold = b
                r2.font.italic = it


def parse_markdown(md_text, doc):
    lines = md_text.split('\n')
    i = 0
    in_yaml = False
    yaml_lines = []
    table_buf = []
    in_code = False

    def flush_table():
        nonlocal table_buf
        if table_buf:
            rows = []
            for ln in table_buf:
                ln = ln.strip()
                if ln.startswith('|'):
                    ln = ln[1:]
                if ln.endswith('|'):
                    ln = ln[:-1]
                rows.append([c.strip() for c in ln.split('|')])
            add_table_md(doc, rows)
            table_buf = []

    while i < len(lines):
        ln = lines[i]

        # YAML frontmatter
        if ln.strip() == '---' and i == 0:
            in_yaml = True
            i += 1
            continue
        if in_yaml and ln.strip() == '---':
            in_yaml = False
            add_yaml_block(doc, yaml_lines)
            yaml_lines = []
            i += 1
            continue
        if in_yaml:
            if ln.strip():
                yaml_lines.append(ln)
            i += 1
            continue

        # Code fence
        if ln.startswith('```'):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            add_para(doc, ln, size=9)
            i += 1
            continue

        # Table
        if ln.strip().startswith('|') and '|' in ln.strip()[1:]:
            table_buf.append(ln)
            i += 1
            continue
        else:
            flush_table()

        # Horizontal rule (non-yaml)
        if ln.strip() == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single')
            bot.set(qn('w:sz'), '4')
            bot.set(qn('w:space'), '1')
            bot.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bot)
            pPr.append(pBdr)
            i += 1
            continue

        # Headers
        if ln.startswith('# '):
            add_h1(doc, ln[2:].strip())
        elif ln.startswith('## '):
            add_h2(doc, ln[3:].strip())
        elif ln.startswith('### '):
            add_h3(doc, ln[4:].strip())
        elif ln.startswith('#### '):
            add_h3(doc, ln[5:].strip())
        elif ln.startswith('> '):
            add_para(doc, ln[2:].strip(), size=10, italic=True, color=CINZA_TXT, quote=True)
        elif ln.startswith('>'):
            add_para(doc, ln[1:].strip(), size=10, italic=True, color=CINZA_TXT, quote=True)
        elif ln.strip().startswith('- ') or ln.strip().startswith('* '):
            add_bullet(doc, ln.strip()[2:])
        elif re.match(r'^\d+\.\s', ln.strip()):
            add_bullet(doc, re.sub(r'^\d+\.\s', '', ln.strip()))
        elif ln.strip() == '':
            pass
        else:
            add_para(doc, ln.strip())

        i += 1

    flush_table()


def convert(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md = f.read()
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    parse_markdown(md, doc)
    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    SRC = Path('C:/Users/leona/orcamentos-openclaw/base/pacotes/arthen-arboris')
    OUT = Path('C:/Users/leona/orcamentos/parametricos/arthen-arboris')
    OUT.mkdir(parents=True, exist_ok=True)
    files = [
        ('PREMISSAS-ORIGEM.md', 'arthen-arboris-PREMISSAS-ORIGEM-v3.docx'),
        ('JUSTIFICATIVA-ITENS-ACIMA-DA-MEDIA.md', 'arthen-arboris-JUSTIFICATIVA-ITENS-ACIMA-MEDIA-v3.docx'),
    ]
    for src, dst in files:
        out = convert(SRC / src, OUT / dst)
        import os
        print(f'OK {out} ({os.path.getsize(out)/1024:.0f} KB)')
