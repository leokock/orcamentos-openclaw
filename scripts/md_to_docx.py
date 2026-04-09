#!/usr/bin/env python3
"""Convert markdown to docx using python-docx."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt

md_path = r'G:\Drives compartilhados\03 CTN Projetos\2. Projetos em Andamento\_Parametrico_IA\arthen-arboris\Memorial-Justificativo-Arthen-Arboris.md'
docx_path = r'G:\Drives compartilhados\03 CTN Projetos\2. Projetos em Andamento\_Parametrico_IA\arthen-arboris\Memorial-Justificativo-Arthen-Arboris.docx'

with open(md_path, encoding='utf-8') as f:
    content = f.read()

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for line in content.split('\n'):
    line = line.rstrip()
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    elif line.startswith('---'):
        continue
    elif line.startswith('|') and '|' in line[1:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if all(set(c) <= set('- :') for c in cells):
            continue
        p = doc.add_paragraph()
        run = p.add_run('  |  '.join(cells))
        run.font.size = Pt(10)
        run.font.name = 'Consolas'
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.strip():
        clean = line.replace('**', '').replace('`', '')
        doc.add_paragraph(clean)

doc.save(docx_path)
print(f'Saved: {docx_path}')
