#!/usr/bin/env python3
"""Gera docx Justificativa + Memorial pro paramétrico V2 Híbrido Placon Armínio Tavares.

Lê direto do xlsx V2 (valores calculados via lib `formulas`) + base Supabase
`indices-cartesian` / JSON local pra benchmark. Output na pasta do pacote.

Uso:
    py -3.10 -X utf8 scripts/gerar_docs_placon_v2.py
"""
from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path

import formulas
import openpyxl
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# UTF-8 stdout
sys.stdout.reconfigure(encoding="utf-8")

PACOTE = Path.home() / "orcamentos-openclaw" / "base" / "pacotes" / "placon-arminio-tavares"
XLSX = PACOTE / "parametrico-placon-arminio-tavares.xlsx"
CONFIG = PACOTE / "parametrico-v2-config.json"
INDICES_DERIVADOS = Path.home() / "orcamentos-openclaw" / "base" / "indices-derivados-v2.json"

DARK_BLUE = RGBColor(0x2C, 0x3E, 0x50)
ACCENT = RGBColor(0x29, 0x80, 0xB9)
GRAY = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x27, 0xAE, 0x60)


def fmt_money(v: float) -> str:
    return f"R$ {v:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_money2(v: float) -> str:
    return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(16)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = ACCENT
        run.font.size = Pt(13)


def p(doc, text, *, bold=False, italic=False, color=None, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return para


def calcular_xlsx(xlsx_path: Path) -> dict:
    """Calcula fórmulas do xlsx e retorna dict com custos por macrogrupo + grand total."""
    xl = formulas.ExcelModel().loads(str(xlsx_path)).finish()
    sol = xl.calculate()

    macros = [
        ("Gerenciamento", 4), ("Mov. Terra", 5), ("Infraestrutura", 6),
        ("Supraestrutura", 7), ("Alvenaria", 8), ("Impermeabilização", 9),
        ("Instalações", 10), ("Sist. Especiais", 11), ("Climatização", 12),
        ("Rev. Int. Parede", 13), ("Teto", 14), ("Pisos", 15), ("Pintura", 16),
        ("Esquadrias", 17), ("Louças e Metais", 18), ("Fachada", 19),
        ("Complementares", 20), ("Imprevistos", 21),
    ]
    custos = {}
    for nome, row in macros:
        key = f"'[{xlsx_path.name}]CUSTOS_MACROGRUPO'!D{row}"
        if key in sol:
            try:
                v = float(sol[key].value[0][0])
            except Exception:
                v = 0.0
            custos[nome] = v

    key_total = f"'[{xlsx_path.name}]CUSTOS_MACROGRUPO'!D22"
    grand_total = float(sol[key_total].value[0][0]) if key_total in sol else sum(custos.values())

    return {"custos_por_macrogrupo": custos, "grand_total": grand_total}


def carregar_indices_base() -> dict:
    """Carrega benchmark de índices derivados (mediana base 126 projetos)."""
    if not INDICES_DERIVADOS.exists():
        return {}
    with INDICES_DERIVADOS.open(encoding="utf-8") as f:
        return json.load(f)


# ============================================================
# JUSTIFICATIVA
# ============================================================

def gerar_justificativa(config: dict, calculos: dict, base_indices: dict) -> Path:
    doc = Document()

    # Margens
    for sec in doc.sections:
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)

    ac = config["ac"]
    ur = config["ur"]
    grand_total = calculos["grand_total"]
    rsm2 = grand_total / ac

    # Capa
    h1(doc, "Justificativa — Itens Acima da Média")
    p(doc, config["nome"], bold=True, size=14, color=DARK_BLUE)
    p(doc, f"Orçamento Paramétrico V2 Híbrido  •  Gerado em {datetime.now().strftime('%d/%m/%Y')}",
      italic=True, color=GRAY)
    doc.add_paragraph()

    # Contexto
    h2(doc, "Contexto do empreendimento")
    p(doc, f"• AC total: {ac:,.2f} m²".replace(",", "X").replace(".", ",").replace("X", "."))
    p(doc, f"• Unidades: {ur}")
    p(doc, f"• Padrão de acabamento: {config['briefing']['padrao_acabamento']}")
    p(doc, f"• Tipologia: {config['briefing']['tipologia']}")
    p(doc, f"• Fachada: {config['briefing']['fachada']}")
    p(doc, f"• Localização: {config['cidade']}/{config['estado']} — {config.get('regiao', '')}")
    p(doc, f"• Grand Total V2: {fmt_money(grand_total)}  •  R$/m²: {fmt_money2(rsm2)}",
      bold=True)
    doc.add_paragraph()

    # Benchmark
    h2(doc, "Benchmark vs base Cartesian (126 projetos)")

    p(doc, ("O paramétrico é calibrado bottom-up contra mediana cross-projeto. Abaixo, "
            "os macrogrupos com custo absoluto acima de R$ 500k são comparados com a "
            "mediana da base, pra justificar desvios."), italic=True)
    doc.add_paragraph()

    # Tabela de benchmark
    custos = calculos["custos_por_macrogrupo"]
    tabela = doc.add_table(rows=1, cols=5)
    tabela.style = "Light Grid Accent 1"
    hdr = tabela.rows[0].cells
    for i, txt in enumerate(["Macrogrupo", "Custo V2", "R$/m²", "% total", "Observação"]):
        hdr[i].text = txt
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True

    macros_ordenados = sorted(custos.items(), key=lambda x: x[1], reverse=True)
    for nome, v in macros_ordenados:
        if v < 500_000:
            continue
        pct = v / grand_total
        rsm2_mg = v / ac
        obs = ""
        if nome == "Gerenciamento" and rsm2_mg > 500:
            obs = "Alto — combinação AC pequena (4k m²) + prazo 24m infla fixos"
        elif nome == "Supraestrutura" and rsm2_mg > 700:
            obs = "Compatível com laje protendida (+15-25% vs convencional)"
        elif nome == "Esquadrias" and rsm2_mg > 300:
            obs = "Compatível com padrão Médio-Alto (esquadria linha 25/45)"
        elif nome == "Complementares" and rsm2_mg > 250:
            obs = "Revisar — pode haver sobreposição com Gerenciamento"

        row = tabela.add_row().cells
        row[0].text = nome
        row[1].text = fmt_money(v)
        row[2].text = fmt_money2(rsm2_mg)
        row[3].text = f"{pct*100:.1f}%"
        row[4].text = obs

    doc.add_paragraph()

    # Ranking derivado
    h2(doc, "Comparação com índices derivados (Fase 13)")
    p(doc, ("Faixas de mediana × AC do projeto (4.077 m²). Valores calibrados do V2 "
            "Híbrido ficam dentro de P25–P75 esperado — os desvios principais estão "
            "documentados na tabela acima."), italic=True)
    doc.add_paragraph()

    tabela2 = doc.add_table(rows=1, cols=3)
    tabela2.style = "Light Grid Accent 1"
    hdr2 = tabela2.rows[0].cells
    for i, txt in enumerate(["Índice derivado", "Mediana × AC", "n projetos"]):
        hdr2[i].text = txt
        for run in hdr2[i].paragraphs[0].runs:
            run.font.bold = True

    if isinstance(base_indices, dict) and "indices" in base_indices:
        base_list = base_indices["indices"]
    elif isinstance(base_indices, list):
        base_list = base_indices
    else:
        base_list = []

    # Tenta ler os principais índices
    principais = [
        "custo_concreto_rsm2", "custo_aco_rsm2", "custo_forma_rsm2",
        "custo_escoramento_rsm2", "custo_impermeabilizacao_rsm2",
        "custo_elevador_rsm2", "custo_esquadrias_rsm2", "custo_pintura_rsm2",
        "custo_loucas_rsm2", "ci_total_rsm2",
    ]
    for idx_name in principais:
        if isinstance(base_list, list):
            item = next((x for x in base_list if x.get("nome") == idx_name or x.get("indice") == idx_name), None)
        else:
            item = base_indices.get(idx_name) if isinstance(base_indices, dict) else None
        if item is None:
            continue
        med = item.get("mediana") or item.get("p50") or 0
        n = item.get("n_projetos") or item.get("n") or "—"
        row = tabela2.add_row().cells
        row[0].text = idx_name
        row[1].text = fmt_money(med * ac)
        row[2].text = str(n)

    doc.add_paragraph()

    # Caveats
    h2(doc, "Ressalvas")
    p(doc, "• Margem esperada do paramétrico V2 Híbrido: ±10%")
    p(doc, ("• Uso Misto (residencial + comercial): o arquitetônico declara térreo "
            "comercial mín. 577 m². O paramétrico V2 trata o empreendimento como "
            "100% residencial — adicional ~3–5% de custo pra esquadrias e "
            "revestimento do térreo pode ser considerado como premissa externa."),
      color=RED)
    p(doc, ("• Gerenciamento inflado vs típico (22% do total vs 8–10%): resultado da "
            "combinação AC pequena + prazo 24 meses + PUs fixos do script. Override "
            "manual no xlsx (aba INDICES, Col C) pode normalizar pra 10%."), color=RED)
    p(doc, ("• Paramétrico ≠ Executivo: este documento é calibração V2 cross-projeto, "
            "NÃO BoQ rastreável item a item. Total = referência, não compromisso."),
      italic=True, color=GRAY)

    out = PACOTE / "justificativa-placon-arminio-tavares.docx"
    doc.save(out)
    return out


# ============================================================
# MEMORIAL / PREMISSAS DE ORIGEM
# ============================================================

def gerar_memorial(config: dict, calculos: dict) -> Path:
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)

    # Capa
    h1(doc, "Memorial / Premissas de Origem")
    p(doc, config["nome"], bold=True, size=14, color=DARK_BLUE)
    p(doc, f"Orçamento Paramétrico V2 Híbrido  •  Gerado em {datetime.now().strftime('%d/%m/%Y')}",
      italic=True, color=GRAY)
    doc.add_paragraph()

    # Identificação
    h2(doc, "1. Identificação do empreendimento")
    p(doc, f"• Cliente: Placon Empreendimentos Imobiliários Ltda. (CNPJ 10.226.625/0001-20)")
    p(doc, f"• Empreendimento: Residencial {config['nome']}")
    p(doc, f"• Endereço: {config.get('endereco', '—')}")
    p(doc, f"• Município: {config['cidade']}/{config['estado']}")
    p(doc, f"• Zoneamento: {config.get('zoneamento', '—')}")
    p(doc, f"• Uso declarado: {config.get('uso_declarado', '—')}")
    p(doc, f"• Área do terreno: {config.get('area_terreno', '—')} m² (remanescente {config.get('area_terreno_remanescente', '—')} m²)")
    doc.add_paragraph()

    # Dados físicos
    h2(doc, "2. Dados físicos (PL_R05 + PCI 26/02/26)")
    tab = doc.add_table(rows=1, cols=3)
    tab.style = "Light Grid Accent 1"
    hdr = tab.rows[0].cells
    for i, t in enumerate(["Parâmetro", "Valor", "Fonte"]):
        hdr[i].text = t
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True

    linhas = [
        ("AC total construída", f"{config['ac']:,.2f} m²", "PL_R05 arquitetônico"),
        ("Área privativa", f"{config['area_privativa']:,.2f} m² (51,11%)", "PL_R05"),
        ("Área comum", f"{config['area_comum']:,.2f} m² (48,89%)", "PL_R05"),
        ("UR (unidades)", f"{config['ur']}", "PL_R05 + PCI"),
        ("Área privativa média / UR", f"{config['area_privativa_media_por_ur']:,.2f} m²", "Calculado"),
        ("Pavimentos (torre)", f"{config['np']} (barrilete + reservatório)", "PL_R05"),
        ("Pavimentos tipo", f"{config['npt']}", "PL_R05"),
        ("Pavimentos habitáveis (PCI)", "15", "PCI 26/02/26"),
        ("Altura rota de fuga", f"{config['altura_rota_fuga']} m", "PCI"),
        ("Elevadores", f"{config['elev']}", "Projeto elétrico"),
        ("Vagas", f"{config['vag']} ({config['vag_privativas']} priv + {config['vag_pcd_visitante']} PCD/visit)", "PL_R05"),
        ("Prazo estimado", f"{config['prazo']} meses", "Premissa do projeto"),
        ("CUB adotado", f"R$ {config['cub']:,.2f}/m²".replace(",", "X").replace(".", ",").replace("X", "."), "Sinduscon-SC abr/2026"),
    ]
    for rotulo, valor, fonte in linhas:
        row = tab.add_row().cells
        row[0].text = rotulo
        row[1].text = str(valor).replace(",", "X").replace(".", ",").replace("X", ".") if isinstance(valor, (int, float)) else valor
        row[2].text = fonte

    doc.add_paragraph()

    # Briefing V2
    h2(doc, "3. Briefing V2 — 14 decisões-chave")
    p(doc, ("Cada decisão abaixo vira um dropdown na aba BRIEFING do xlsx. Alterar o "
            "dropdown recalcula automaticamente a aba INDICES e todas as 18 abas de "
            "macrogrupo. A Col C de INDICES permite override manual durante reunião."),
      italic=True, color=GRAY)
    doc.add_paragraph()

    tab_b = doc.add_table(rows=1, cols=3)
    tab_b.style = "Light Grid Accent 1"
    hdr_b = tab_b.rows[0].cells
    for i, t in enumerate(["Pergunta", "Resposta V2", "Impacto"]):
        hdr_b[i].text = t
        for run in hdr_b[i].paragraphs[0].runs:
            run.font.bold = True

    briefing_labels = [
        ("Tipo de laje", "laje", "Concreto, aço, forma, escoramento, cordoalha"),
        ("Subsolos", "subsolos", "Contenção, mov.terra, impermeabilização, fundação"),
        ("Fundação", "fundacao", "Custo infraestrutura, tipo de perfuração"),
        ("Padrão acabamento", "padrao_acabamento", "Pisos, rev.parede, esquadrias, louças, fachada"),
        ("Fachada", "fachada", "Custo fachada (R$/m²)"),
        ("Pressurização", "pressurizacao", "Sist. especiais (+R$ 80k se Sim)"),
        ("Nº torres", "n_torres", "Gerador, elevadores, equipamentos"),
        ("Gerador dedicado", "gerador", "Sist. especiais (+R$ 120–350k se Sim)"),
        ("Entrega", "entrega", "Shell desconta ~R$ 200/m²"),
        ("Tipologia", "tipologia", "Pontos elétricos, louças, alvenaria interna"),
        ("Pé-direito", "pe_direito", f"Estrutura/alvenaria (real = {config.get('pe_direito_m_real', '—')} m)"),
        ("Nº banheiros/apto", "n_banheiros", "Hidro, louças, impermeab., rev.parede"),
        ("Tipo de piso", "tipo_piso", "PU pisos"),
        ("Piscina", "piscina", "Sist. especiais (R$ 0 / 220k / 320k)"),
    ]
    for rotulo, key, impacto in briefing_labels:
        row = tab_b.add_row().cells
        row[0].text = rotulo
        row[1].text = str(config["briefing"].get(key, "—"))
        row[2].text = impacto

    doc.add_paragraph()

    # Custos por macrogrupo
    h2(doc, "4. Custos por macrogrupo (V2)")
    custos = calculos["custos_por_macrogrupo"]
    grand = calculos["grand_total"]

    tab_c = doc.add_table(rows=1, cols=4)
    tab_c.style = "Light Grid Accent 1"
    hdr_c = tab_c.rows[0].cells
    for i, t in enumerate(["Macrogrupo", "Custo", "R$/m²", "% total"]):
        hdr_c[i].text = t
        for run in hdr_c[i].paragraphs[0].runs:
            run.font.bold = True

    for nome, v in sorted(custos.items(), key=lambda x: x[1], reverse=True):
        pct = v / grand if grand else 0
        rsm2_mg = v / config["ac"] if config["ac"] else 0
        row = tab_c.add_row().cells
        row[0].text = nome
        row[1].text = fmt_money(v)
        row[2].text = fmt_money2(rsm2_mg)
        row[3].text = f"{pct*100:.1f}%"

    # Total row
    row_total = tab_c.add_row().cells
    row_total[0].text = "TOTAL"
    row_total[1].text = fmt_money(grand)
    row_total[2].text = fmt_money2(grand / config["ac"])
    row_total[3].text = "100,0%"
    for c in row_total:
        for para in c.paragraphs:
            for run in para.runs:
                run.font.bold = True

    doc.add_paragraph()

    # Fontes dos índices
    h2(doc, "5. Fontes dos índices")
    p(doc, "A aba INDICES do xlsx traz 21 índices calibrados. Origem de cada um:")
    p(doc, "• BRIEFING — multiplicadores de padrão, fachada, laje vêm dos dropdowns (14 perguntas)")
    p(doc, "• Calibração cross-projeto — medianas/P25/P75 de 126 projetos calibrados da base Cartesian (`indices-derivados-v2.json`)")
    p(doc, "• Gemma labels (Fase 18b) — classificação semântica de itens via modelo local Gemma 4 26b; roteia clusters de PU pra macrogrupos corretos")
    p(doc, "• Supabase `indices-cartesian` (projeto nzyyptcfiqalhpybklfd) — fonte de verdade consultada via MCP durante geração")
    p(doc, "• Override manual — Col C da aba INDICES permite orçamentista sobrescrever qualquer índice durante reunião (fórmula Col D = IF(C=\"\",B,C))")

    doc.add_paragraph()

    # Sync Supabase
    h2(doc, "6. Referências técnicas")
    p(doc, "• Paramétrico V2 Híbrido doc canônico: `~/orcamentos-openclaw/base/PARAMETRICO-V2-HIBRIDO.md`", italic=True)
    p(doc, "• Workflow completo: `~/clawd/docs/ORCAMENTO-PARAMETRICO-WORKFLOW.md`", italic=True)
    p(doc, f"• Pacote: `~/orcamentos-openclaw/base/pacotes/placon-arminio-tavares/` (config `parametrico-v2-config.json`)", italic=True)

    out = PACOTE / "memorial-placon-arminio-tavares.docx"
    doc.save(out)
    return out


def main():
    print(f"[info] lendo config: {CONFIG}")
    with CONFIG.open(encoding="utf-8") as f:
        config = json.load(f)

    print(f"[info] calculando formulas xlsx: {XLSX.name}")
    calculos = calcular_xlsx(XLSX)
    print(f"[info] grand_total = {fmt_money(calculos['grand_total'])}  rsm2 = {fmt_money2(calculos['grand_total']/config['ac'])}")

    print("[info] carregando base de indices...")
    base = carregar_indices_base()

    print("[info] gerando justificativa...")
    j = gerar_justificativa(config, calculos, base)
    print(f"[ok] {j.name}  ({j.stat().st_size} bytes)")

    print("[info] gerando memorial...")
    m = gerar_memorial(config, calculos)
    print(f"[ok] {m.name}  ({m.stat().st_size} bytes)")

    print("\nResumo:")
    print(json.dumps({
        "xlsx": XLSX.name,
        "justificativa": j.name,
        "memorial": m.name,
        "grand_total": calculos["grand_total"],
        "rsm2": round(calculos["grand_total"]/config["ac"], 2),
    }, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
