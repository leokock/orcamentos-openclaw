#!/usr/bin/env python3
"""
Script para gerar briefing Word do Thozen Electra
Organizado por pavimento (estrutura da planilha Excel)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_color(cell, color_hex):
    """Adiciona cor de fundo à célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def criar_tabela_pavimento(doc, pavimento, descricao, un, qtd, observacao):
    """Cria tabela de quantitativos por pavimento"""
    # Criar tabela com 5 colunas
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Cabeçalho
    hdr_cells = table.rows[0].cells
    headers = ['Pavimento', 'Descrição', 'UN', 'QTD', 'Observação']
    
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Aplicar cor azul
        set_cell_color(hdr_cells[i], '4472C4')
        # Formatar texto
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Adicionar linhas de dados
    for i in range(len(pavimento)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(pavimento[i])
        row_cells[1].text = str(descricao[i])
        row_cells[2].text = str(un[i])
        row_cells[3].text = str(qtd[i])
        row_cells[4].text = str(observacao[i])
    
    return table

def main():
    # Criar documento
    doc = Document()
    
    # Configurar margens (2,5cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Definir estilos
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    
    # =====================
    # CAPA
    # =====================
    title = doc.add_paragraph()
    title_run = title.add_run('BRIEFING TÉCNICO\nTHOZEN ELECTRA\nQUANTITATIVOS POR PAVIMENTO')
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_before = Pt(144)
    
    doc.add_paragraph()
    
    info = doc.add_paragraph('Projeto: Edifício Residencial Thozen Electra', style='Normal')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = doc.add_paragraph('Endereço: Rua Rubens Alves', style='Normal')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = doc.add_paragraph('Data: 20/03/2026', style='Normal')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = doc.add_paragraph('Revisão: R00', style='Normal')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # =====================
    # SUMÁRIO EXECUTIVO
    # =====================
    h1 = doc.add_heading('SUMÁRIO EXECUTIVO', level=1)
    h1.runs[0].font.color.rgb = RGBColor(68, 114, 196)
    
    doc.add_paragraph(
        'Este briefing apresenta os quantitativos técnicos do projeto Thozen Electra organizados por pavimento, '
        'conforme estrutura da planilha Excel de quantitativos.'
    )
    
    doc.add_paragraph('Estrutura do edifício:', style='List Bullet')
    doc.add_paragraph('Térreo (1º pavimento)', style='List Bullet 2')
    doc.add_paragraph('G1 a G5 (2º ao 6º pavimento - garagens)', style='List Bullet 2')
    doc.add_paragraph('Lazer (7º pavimento)', style='List Bullet 2')
    doc.add_paragraph('Tipos (8º a 31º pavimento - 24 pavimentos)', style='List Bullet 2')
    doc.add_paragraph('Residência + Cobertura (32º pavimento)', style='List Bullet 2')
    doc.add_paragraph('Casa de Máquinas', style='List Bullet 2')
    
    p = doc.add_paragraph()
    p.add_run('Total de pavimentos: ').font.bold = True
    p.add_run('35')
    
    doc.add_paragraph('Disciplinas abrangidas:')
    disciplinas = [
        '1. Estrutura', '2. Alvenaria', '3. Hidráulico', '4. Sanitário',
        '5. PCI Civil', '6. PCI Elétrico', '7. Elétrico', '8. Telefônico',
        '9. SPDA', '10. Ventilação', '11. Exaustão', '12. Ar-Condicionado',
        '13. Esquadrias'
    ]
    for disc in disciplinas:
        doc.add_paragraph(disc, style='List Bullet')
    
    doc.add_page_break()
    
    # =====================
    # SEÇÃO 1: ESTRUTURA
    # =====================
    h1 = doc.add_heading('1. ESTRUTURA', level=1)
    h1.runs[0].font.color.rgb = RGBColor(68, 114, 196)
    
    p = doc.add_paragraph()
    p.add_run('Status: ').font.bold = True
    run = p.add_run('✅ Quantitativos extraídos do IFC')
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    p = doc.add_paragraph()
    p.add_run('Fonte: ').font.bold = True
    p.add_run('IFC R26 - 1203 THOZEN - RUBENS ALVES - BLOCOS+RAMPAS DE ACESSO')
    
    doc.add_heading('1.1 Quantitativos por Pavimento', level=2)
    
    # Infraestrutura
    doc.add_paragraph('INFRAESTRUTURA:', style='Heading 3')
    
    pavimento = ['Térreo', 'Térreo', 'Térreo', 'Térreo', 'Térreo']
    descricao = [
        'Vigas baldrame 20×250cm',
        'Vigas baldrame 14×184cm',
        'Vigas baldrame 14×164cm',
        'Vigas baldrame 14×124cm',
        'Blocos/pranchas fundação'
    ]
    un = ['m³', 'm³', 'm³', 'm³', 'un']
    qtd = ['24,00', '4,62', '15,87', '10,92', '4']
    obs = [
        'Vigas de perímetro',
        'Vigas internas',
        'Vigas internas',
        'Vigas menores',
        'Blocos de apoio'
    ]
    
    criar_tabela_pavimento(doc, pavimento, descricao, un, qtd, obs)
    
    p = doc.add_paragraph()
    p.add_run('Total: ').font.bold = True
    p.add_run('~55 m³ de concreto infraestrutura + 4 blocos')
    
    doc.add_paragraph()
    
    # Supraestrutura - Pilares
    doc.add_paragraph('SUPRAESTRUTURA - PILARES:', style='Heading 3')
    
    pavimento_pil = ['Térreo', 'G1-G5', 'Lazer', '8º-31º Tipo', 'Res/Cob', 'Casa Máq.']
    desc_pil = [
        'Pilares diversos',
        'Pilares (5 pavimentos)',
        'Pilares',
        'Pilares (×24 pavimentos)',
        'Pilares',
        'Pilares'
    ]
    un_pil = ['m³', 'm³', 'm³', 'm³', 'm³', 'm³']
    qtd_pil = ['~200', '~700', '~130', '~1.100', '~70', '~15']
    obs_pil = [
        'Térreo + rampas (72 un)',
        '~49-51 pilares/pav',
        '46 pilares',
        '~39-42 pilares/pav',
        '25 pilares',
        '5 pilares'
    ]
    
    criar_tabela_pavimento(doc, pavimento_pil, desc_pil, un_pil, qtd_pil, obs_pil)
    
    p = doc.add_paragraph()
    p.add_run('Total: ').font.bold = True
    p.add_run('2.964 m³ de concreto em pilares (1.531 elementos)')
    
    doc.add_paragraph()
    
    # Supraestrutura - Vigas
    doc.add_paragraph('SUPRAESTRUTURA - VIGAS:', style='Heading 3')
    
    pavimento_vig = ['Térreo', 'G1-G5', 'Lazer', '8º-31º Tipo', 'Res/Cob', 'Casa Máq.']
    desc_vig = [
        'Vigas diversas',
        'Vigas (5 pavimentos)',
        'Vigas (estrutura complexa)',
        'Vigas (×24 pavimentos)',
        'Vigas cobertura',
        'Vigas barrilete'
    ]
    un_vig = ['m³', 'm³', 'm³', 'm³', 'm³', 'm³']
    qtd_vig = ['~400', '~350', '~450', '~2.600', '~250', '~30']
    obs_vig = [
        '~170 vigas',
        '~60-110 vigas/pav',
        '~190 vigas',
        '~100-110 vigas/pav',
        '~105 vigas',
        '~13 vigas'
    ]
    
    criar_tabela_pavimento(doc, pavimento_vig, desc_vig, un_vig, qtd_vig, obs_vig)
    
    p = doc.add_paragraph()
    p.add_run('Total: ').font.bold = True
    p.add_run('2.406 m³ de concreto em vigas (3.531 elementos)')
    
    doc.add_paragraph()
    
    # Supraestrutura - Lajes
    doc.add_paragraph('SUPRAESTRUTURA - LAJES:', style='Heading 3')
    
    pavimento_laj = ['Térreo', 'G1-G5', 'Lazer', '8º-31º Tipo', 'Res/Cob', 'Casa Máq.']
    desc_laj = [
        'Lajes e<|end_of_text|>