#!/usr/bin/env python3
"""
Script para integrar informações do cliente e empreendimento no briefing do Thozen Electra.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

def add_page_number(section):
    """Adiciona numeração de página ao rodapé"""
    footer = section.footer
    footer.is_linked_to_previous = False
    
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar número de página
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    
    run.font.size = Pt(9)

def add_footer_text(section):
    """Adiciona texto do rodapé em todas as seções"""
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # Usar o primeiro parágrafo ou criar um novo
    if footer.paragraphs:
        p = footer.paragraphs[0]
        p.text = ""  # Limpar conteúdo existente
    else:
        p = footer.add_paragraph()
    
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run("Cliente: Thozen Construtora | Projeto: Electra Towers | Data: 20/03/2026")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)

def main():
    input_file = "output/Briefing-Thozen-Electra-Completo-Por-Pavimento.docx"
    output_file = "output/Briefing-Thozen-Electra-Completo-v2.docx"
    
    print(f"📄 Carregando documento: {input_file}")
    doc = Document(input_file)
    
    # ===== 1. ADICIONAR SEÇÃO "INFORMAÇÕES DO PROJETO" =====
    print("\n✏️  Adicionando seção 'Informações do Projeto'...")
    
    # Encontrar posição do "Sumário Executivo" (assumindo que está logo após a capa)
    insert_index = 0
    for i, paragraph in enumerate(doc.paragraphs):
        if "SUMÁRIO EXECUTIVO" in paragraph.text.upper() or "Sumário Executivo" in paragraph.text:
            insert_index = i
            break
    
    # Inserir título da seção
    p = doc.paragraphs[insert_index].insert_paragraph_before()
    p.add_run("1. INFORMAÇÕES DO PROJETO").bold = True
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    # 1.1 Cliente
    p = doc.paragraphs[insert_index + 1].insert_paragraph_before()
    p.add_run("1.1 Cliente").bold = True
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    
    cliente_info = [
        "• Cliente: Thozen Construtora e Incorporadora",
        "• Localização: Av. Tab. Nomi Jacó Cruz, 71 - Porto Belo - SC, 88210-000",
        "• Site: https://thozenconstrutora.com.br",
        "• Cliente recorrente: Sim",
        "• Projetos anteriores: Mirador de Alicante (Cartesian)",
        "• Contato: Nicholas (Coordenador de obras) - engenharia@thozen.com.br"
    ]
    
    for info in cliente_info:
        p = doc.paragraphs[insert_index + 2].insert_paragraph_before()
        run = p.add_run(info)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.paragraph_format.line_spacing = 1.15
    
    # 1.2 Empreendimento
    p = doc.paragraphs[insert_index + 8].insert_paragraph_before()
    p.add_run("1.2 Empreendimento").bold = True
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    
    emp_info = [
        "• Nome: Electra Towers",
        "• Localização: Rua Canoinhas & R. Rubens Alves, Porto Belo - SC, 88210-000",
        "• Área total do empreendimento: 37.893,89 m²",
        "• Estrutura:",
        "  - 2 torres (Torre A + Torre B)",
        "  - Térreo",
        "  - 5 garagens (G1-G5)",
        "  - Lazer (7º pavimento)",
        "  - 24 pavimentos tipo por torre (8º-31º)",
        "  - Casa de Máquinas",
        "  - Reservatório"
    ]
    
    for info in emp_info:
        p = doc.paragraphs[insert_index + 9].insert_paragraph_before()
        run = p.add_run(info)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.paragraph_format.line_spacing = 1.15
    
    # 1.3 Escopo Contratado
    p = doc.paragraphs[insert_index + 20].insert_paragraph_before()
    p.add_run("1.3 Escopo Contratado").bold = True
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    
    escopo_info = [
        "• Modelo para extração de quantitativos",
        "• Orçamento executivo",
        "• Planejamento executivo",
        "• Integração de tempo e custo",
        "• ERP do cliente: Sienge",
        "• Software de planejamento: Prevision",
        "• Desembolso para insumos curva A/B: Sim",
        "• 5D: Não incluído"
    ]
    
    for info in escopo_info:
        p = doc.paragraphs[insert_index + 21].insert_paragraph_before()
        run = p.add_run(info)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.paragraph_format.line_spacing = 1.15
    
    # ===== 2. ATUALIZAR SUMÁRIO EXECUTIVO =====
    print("\n✏️  Atualizando Sumário Executivo...")
    
    # Procurar parágrafos com área e custo/m²
    for paragraph in doc.paragraphs:
        text = paragraph.text
        
        # Atualizar área total
        if "25.000" in text and "m²" in text:
            new_text = text.replace("25.000", "37.893,89")
            paragraph.text = new_text
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
        
        # Atualizar custo/m²
        if "850" in text and "R$" in text and "/m²" in text:
            new_text = text.replace("R$ 850", "R$ 561")
            paragraph.text = new_text
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
        
        # Atualizar estrutura (se houver menção genérica)
        if "estrutura" in text.lower() and len(text) < 200:
            # Adicionar descrição detalhada
            if "2 torres" not in text.lower():
                p_new = paragraph.insert_paragraph_before()
                run = p_new.add_run("Estrutura do empreendimento: 2 torres (Torre A + Torre B), térreo, 5 garagens (G1-G5), lazer (7º pavimento), 24 pavimentos tipo por torre (8º-31º), casa de máquinas e reservatório.")
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                p_new.paragraph_format.line_spacing = 1.15
    
    # ===== 3. ADICIONAR RODAPÉ =====
    print("\n✏️  Adicionando rodapé em todas as páginas...")
    
    for section in doc.sections:
        add_footer_text(section)
    
    # ===== 4. SALVAR DOCUMENTO =====
    print(f"\n💾 Salvando documento atualizado: {output_file}")
    doc.save(output_file)
    
    print("\n✅ Briefing atualizado com sucesso!")
    print(f"\n📊 Resumo das alterações:")
    print("  • Seção 'Informações do Projeto' adicionada (Cliente, Empreendimento, Escopo)")
    print("  • Área total atualizada: 25.000 m² → 37.893,89 m²")
    print("  • Custo/m² atualizado: R$ 850/m² → R$ 561/m²")
    print("  • Rodapé adicionado: Cliente: Thozen Construtora | Projeto: Electra Towers | Data: 20/03/2026")
    print(f"  • Arquivo gerado: {output_file}")

if __name__ == "__main__":
    main()
