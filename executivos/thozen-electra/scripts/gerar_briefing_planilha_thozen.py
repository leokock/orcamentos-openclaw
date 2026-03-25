#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para gerar briefing consolidado em Word (.docx) e planilha detalhada em Excel (.xlsx)
do projeto Thozen Electra com TODOS os quantitativos extraídos das 14 disciplinas.

Data: 2026-03-20
Projeto: Thozen Electra
Cliente: Thozen
Projetista: R. Rubens Alves
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.chart import PieChart, Reference
from datetime import datetime
import os

# ===========================
# DADOS CONSOLIDADOS
# ===========================

PROJETO = {
    "nome": "Thozen Electra",
    "cliente": "Thozen",
    "projetista": "R. Rubens Alves",
    "tipologia": "Edifício residencial vertical",
    "pavimentos": 32,
    "torres": 2,
    "altura_m": 87,
    "area_total_m2": 25000,  # estimado
    "data": "20/03/2026"
}

# Estrutura de disciplinas
DISCIPLINAS = [
    {
        "codigo": "01",
        "nome": "ESTRUTURA",
        "status": "Completo",
        "qtd_itens": 15,
        "custo_estimado": 5280000,
        "incerteza": "±15%",
        "fonte": "IFC R26",
        "principais": [
            ("Concreto (infraestrutura)", "m³", 54.58),
            ("Concreto (pilares)", "m³", 2964.25),
            ("Concreto (vigas)", "m³", 2406.38),
            ("Concreto (lajes)", "m³", 7358.80),
            ("Aço (total estimado)", "ton", 1020)  # 80kg/m³ × 12.784m³
        ],
        "obs": "Estacas não identificadas no IFC — aguardando prancha de fundações"
    },
    {
        "codigo": "02",
        "nome": "ARQUITETURA",
        "status": "Parcial",
        "qtd_itens": 12,
        "custo_estimado": 3800000,
        "incerteza": "±25%",
        "fonte": "IFC R07+R08",
        "principais": [
            ("Espaços (ambientes)", "un", 838),
            ("Portas + Janelas", "un", 4488),
            ("Tipologias de piso", "tipos", 18),
            ("Lajes/Pisos estruturais", "un", 2610)
        ],
        "obs": "Áreas de revestimentos pendentes — necessário DWG"
    },
    {
        "codigo": "03",
        "nome": "ALVENARIA",
        "status": "Parcial",
        "qtd_itens": 8,
        "custo_estimado": 1200000,
        "incerteza": "±30%",
        "fonte": "DXF (5 pavimentos)",
        "principais": [
            ("Comprimento de paredes", "m", 941.21),
            ("Portas (vãos)", "un", 237)
        ],
        "obs": "4 pavimentos de garagem não processados (arquivos >300MB)"
    },
    {
        "codigo": "04",
        "nome": "ESQUADRIAS",
        "status": "Completo",
        "qtd_itens": 25,
        "custo_estimado": 1850000,
        "incerteza": "±12%",
        "fonte": "IFC R07 + JSON",
        "principais": [
            ("Portas de madeira", "un", 2111),
            ("Portas de alumínio", "un", 238),
            ("Portas de elevador", "un", 186),
            ("Janelas basculantes", "un", 552),
            ("Janelas de correr", "un", 889),
            ("Vidros fixos", "un", 17)
        ],
        "obs": "Peles de vidro pendentes — consultar DWG específico"
    },
    {
        "codigo": "05",
        "nome": "HIDRÁULICO",
        "status": "Completo",
        "qtd_itens": 42,
        "custo_estimado": 980000,
        "incerteza": "±10%",
        "fonte": "IFC rev.01",
        "principais": [
            ("Tubulações totais", "m", 291556.56),
            ("PVC Soldável", "m", 193367.40),
            ("CPVC FlowGuard", "m", 62519.86),
            ("PPR PN25", "m", 35669.30),
            ("Conexões", "un", 3116),
            ("Registros/Válvulas", "un", 270),
            ("Reservatórios 7.500L", "un", 2),
            ("Pressurizadores 2CV", "un", 4)
        ],
        "obs": "Louças e metais não identificados no IFC"
    },
    {
        "codigo": "06",
        "nome": "SANITÁRIO",
        "status": "Parcial",
        "qtd_itens": 38,
        "custo_estimado": 820000,
        "incerteza": "±20%",
        "fonte": "IFC rev.01 (10 arquivos)",
        "principais": [
            ("Tubulações PVC", "trechos", 7326),
            ("Tubulações PVC (metragem)", "m", 106680.71),
            ("Conexões e acessórios", "un", 11340),
            ("Caixas de inspeção lodo", "un", 19),
            ("Tanques lodo ativado (ETE)", "un", 12)
        ],
        "obs": "Metragens de tubulação com valores zerados em 90% dos pavimentos — recalcular"
    },
    {
        "codigo": "07",
        "nome": "PCI CIVIL",
        "status": "Preliminar",
        "qtd_itens": 18,
        "custo_estimado": 450000,
        "incerteza": "±40%",
        "fonte": "IFC rev.01 (4 arquivos)",
        "principais": [
            ("Tubulação FG Ø150mm", "m", 67.26),  # subestimado
            ("Abrigos de hidrante", "un", 67),
            ("Extintores PQS 4kg", "un", 133),
            ("Extintores CO2 6kg", "un", 7),
            ("Placas fotoluminescentes", "un", 140),
            ("Suportes de parede", "un", 135)
        ],
        "obs": "Metragem de tubulação subestimada — validar DWGs. Reservatórios e bombas não modelados"
    },
    {
        "codigo": "08",
        "nome": "PCI ELÉTRICO",
        "status": "Completo",
        "qtd_itens": 31,
        "custo_estimado": 580000,
        "incerteza": "±15%",
        "fonte": "IFC rev.01 (9 arquivos)",
        "principais": [
            ("Detectores de fumaça", "un", 494),
            ("Acionadores manuais", "un", 72),
            ("Avisadores audiovisuais", "un", 102),
            ("Centrais de alarme", "un", 3),
            ("Quadros de comando", "un", 40)  # ajustado — IFC reportou 1.189
        ],
        "obs": "Cabos e eletrodutos não extraídos — estimar 15m/ponto"
    },
    {
        "codigo": "09",
        "nome": "ELÉTRICO",
        "status": "Parcial",
        "qtd_itens": 28,
        "custo_estimado": 2100000,
        "incerteza": "±25%",
        "fonte": "IFC rev.01 (9 arquivos)",
        "principais": [
            ("Luminárias", "un", 837),  # sem multiplicador tipo
            ("Luminárias (total estimado)", "un", 4661),  # com tipo ×24
            ("Eletrodutos (trechos)", "trechos", 18967),
            ("Cabos (trechos)", "trechos", 1014)
        ],
        "obs": "Diâmetros de eletrodutos e bitolas de cabos não especificados — consultar DWGs"
    },
    {
        "codigo": "10",
        "nome": "TELEFÔNICO",
        "status": "Completo",
        "qtd_itens": 24,
        "custo_estimado": 320000,
        "incerteza": "±12%",
        "fonte": "IFC rev.01 (9 arquivos)",
        "principais": [
            ("Pontos de dados RJ45", "un", 46),
            ("Pontos de voz RJ11", "un", 44),
            ("Caixas de passagem 4×2", "un", 289),
            ("Caixas de passagem 4×4", "un", 87),
            ("Eletrodutos", "m", 5830),  # sem multiplicador
            ("Eletrodutos (total estimado)", "m", 33400),  # com tipo ×24
            ("Eletrocalhas perfuradas", "m", 33)
        ],
        "obs": "Cabos UTP, racks e patch panels não modelados — consultar memorial"
    },
    {
        "codigo": "11",
        "nome": "SPDA",
        "status": "Estimado",
        "qtd_itens": 16,
        "custo_estimado": 180000,
        "incerteza": "±20%",
        "fonte": "DWG (10 arquivos) — estimativa",
        "principais": [
            ("Captores tipo Franklin", "un", 6),
            ("Descidas (cabo 50mm²)", "m", 354),
            ("Anéis equipotenciais (35mm²)", "m", 600),
            ("Malha aterramento (50mm²)", "m", 1000),
            ("Hastes cobreadas Ø5/8×2,4m", "un", 8),
            ("Conectores bronze", "un", 212)
        ],
        "obs": "Quantitativos baseados em premissas — validar com memorial descritivo"
    },
    {
        "codigo": "12",
        "nome": "VENTILAÇÃO",
        "status": "Completo (DXF)",
        "qtd_itens": 35,
        "custo_estimado": 720000,
        "incerteza": "±12%",
        "fonte": "DXF R05 (30 MB)",
        "principais": [
            ("Ventiladores pressurização", "un", 4),
            ("Exaustores desenfumagem", "un", 10),
            ("Grelhas GR-01", "un", 34),
            ("Difusores DF-01/DF-02", "un", 120),
            ("Dutos (metragem bruta)", "m", 22200),
            ("Inversores Fire Mode", "un", 10),
            ("Central CPS-B1-5-0101", "un", 2)
        ],
        "obs": "Sistema DUAL: Pressurização + Desenfumagem. Vazões pendentes de memorial"
    },
    {
        "codigo": "13",
        "nome": "EXAUSTÃO",
        "status": "Completo (DXF)",
        "qtd_itens": 28,
        "custo_estimado": 580000,
        "incerteza": "±15%",
        "fonte": "DXF R00 (18 MB)",
        "principais": [
            ("Churrasqueiras", "un", 195),
            ("Exaustores TCV 710 (3kW)", "un", 8),
            ("Prumadas verticais", "un", 8),
            ("Coifas inox AISI 304", "un", 195),
            ("Dutos metálicos (estimado)", "m", 1600),
            ("Dampers corta-fogo 90min", "un", 180),
            ("Inversores de frequência", "un", 8)
        ],
        "obs": "Metragem exata de dutos pendente — DXF reportou 7.378m (provavelmente erro)"
    },
    {
        "codigo": "14",
        "nome": "AR-CONDICIONADO",
        "status": "Completo",
        "qtd_itens": 32,
        "custo_estimado": 3200000,
        "incerteza": "±18%",
        "fonte": "DXF R05 (39,9 MB)",
        "principais": [
            ("Evaporadoras", "un", 1000),
            ("Condensadoras", "un", 997),
            ("Potência total", "BTU/h", 9132000),
            ("Potência total", "TR", 761),
            ("Tubulações frigoríficas", "m", 9000),
            ("Linhas de dreno", "m", 6800),
            ("Disjuntores", "un", 1997)
        ],
        "obs": "Tubulações estimadas — extração do DXF pendente"
    }
]

# Calcular totais
TOTAL_CUSTO = sum(d["custo_estimado"] for d in DISCIPLINAS)
TOTAL_ITENS = sum(d["qtd_itens"] for d in DISCIPLINAS)

# ===========================
# FUNÇÕES AUXILIARES
# ===========================

def adicionar_borda_tabela(tabela):
    """Adiciona bordas simples a uma tabela do Word"""
    tbl = tabela._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for borda in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        elemento = OxmlElement(f'w:{borda}')
        elemento.set(qn('w:val'), 'single')
        elemento.set(qn('w:sz'), '4')
        elemento.set(qn('w:color'), '000000')
        tblBorders.append(elemento)
    
    tblPr.append(tblBorders)

def formatar_moeda(valor):
    """Formata valor em R$"""
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def formatar_numero(valor, decimais=2):
    """Formata número com separador de milhares"""
    if decimais == 0:
        return f"{int(valor):,}".replace(",", ".")
    else:
        return f"{valor:,.{decimais}f}".replace(",", "X").replace(".", ",").replace("X", ".")

# ===========================
# GERAÇÃO DO WORD
# ===========================

def gerar_briefing_word():
    """Gera briefing consolidado em Word"""
    print("\n🔄 Gerando briefing Word...")
    
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # =============================
    # CAPA
    # =============================
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("BRIEFING TÉCNICO\nTHOZEN ELECTRA")
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Orçamento Executivo - Todas as Disciplinas")
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(68, 68, 68)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Data: {PROJETO['data']}\nCartesian Engenharia")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # =============================
    # SUMÁRIO EXECUTIVO
    # =============================
    secao = doc.add_heading('Sumário Executivo', 1)
    secao.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    # Resumo do projeto
    doc.add_heading('Caracterização do Projeto', 2)
    tabela = doc.add_table(rows=7, cols=2)
    tabela.style = 'Light Grid Accent 1'
    adicionar_borda_tabela(tabela)
    
    dados = [
        ("Nome do Projeto", PROJETO["nome"]),
        ("Cliente", PROJETO["cliente"]),
        ("Projetista", PROJETO["projetista"]),
        ("Tipologia", PROJETO["tipologia"]),
        ("Pavimentos", f"{PROJETO['pavimentos']} pavimentos | {PROJETO['torres']} torres"),
        ("Altura Estimada", f"{PROJETO['altura_m']} metros"),
        ("Área Total Estimada", f"{formatar_numero(PROJETO['area_total_m2'], 0)} m²")
    ]
    
    for i, (campo, valor) in enumerate(dados):
        tabela.rows[i].cells[0].text = campo
        tabela.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        tabela.rows[i].cells[1].text = str(valor)
    
    doc.add_paragraph()
    
    # Status da extração
    doc.add_heading('Status da Extração por Disciplina', 2)
    tabela = doc.add_table(rows=len(DISCIPLINAS)+1, cols=4)
    tabela.style = 'Light Grid Accent 1'
    adicionar_borda_tabela(tabela)
    
    # Cabeçalho
    cabecalho = ['Disciplina', 'Status', 'Qtd Itens', 'Fonte']
    for i, texto in enumerate(cabecalho):
        cell = tabela.rows[0].cells[i]
        cell.text = texto
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dados
    for i, disc in enumerate(DISCIPLINAS, start=1):
        tabela.rows[i].cells[0].text = f"{disc['codigo']} - {disc['nome']}"
        tabela.rows[i].cells[1].text = disc['status']
        tabela.rows[i].cells[2].text = str(disc['qtd_itens'])
        tabela.rows[i].cells[3].text = disc['fonte']
    
    doc.add_paragraph()
    
    # Custos consolidados
    doc.add_heading('Custos Consolidados', 2)
    p = doc.add_paragraph()
    p.add_run(f"• Custo Total Estimado: ").font.bold = True
    p.add_run(f"{formatar_moeda(TOTAL_CUSTO)}\n")
    p.add_run(f"• Total de Itens Orçados: ").font.bold = True
    p.add_run(f"{TOTAL_ITENS} itens\n")
    p.add_run(f"• Custo por m²: ").font.bold = True
    custo_m2 = TOTAL_CUSTO / PROJETO['area_total_m2']
    p.add_run(f"{formatar_moeda(custo_m2)}/m²")
    
    doc.add_paragraph()
    
    # Principais descobertas
    doc.add_heading('Principais Descobertas', 2)
    descobertas = [
        "Sistema DUAL de ventilação: Pressurização de escadas + Desenfumagem de corredores",
        "Total de 195 churrasqueiras no pavimento Lazer com sistema de exaustão independente",
        "~1.000 evaporadoras de ar-condicionado (splits individuais por ambiente)",
        "Sistema de tratamento de efluentes (ETE) com lodo ativado no térreo",
        "Pavimento Tipo se repete 24 vezes (8º ao 31º andar)",
        "4 pavimentos de garagem não processados em Alvenaria (arquivos >300MB)",
        "Metragens de tubulação sanitária precisam ser recalculadas (90% com valores zerados)"
    ]
    
    for descoberta in descobertas:
        doc.add_paragraph(f"• {descoberta}", style='List Bullet')
    
    doc.add_page_break()
    
    # =============================
    # SEÇÕES POR DISCIPLINA
    # =============================
    for disc in DISCIPLINAS:
        doc.add_heading(f'{disc["codigo"]}. {disc["nome"]}', 1)
        
        # Resumo
        p = doc.add_paragraph()
        p.add_run("Status: ").font.bold = True
        p.add_run(f"{disc['status']}\n")
        p.add_run("Fonte de dados: ").font.bold = True
        p.add_run(f"{disc['fonte']}\n")
        p.add_run("Quantidade de itens: ").font.bold = True
        p.add_run(f"{disc['qtd_itens']}\n")
        p.add_run("Custo estimado: ").font.bold = True
        p.add_run(f"{formatar_moeda(disc['custo_estimado'])} ({disc['incerteza']})")
        
        doc.add_paragraph()
        
        # Tabela de quantitativos
        doc.add_heading('Quantitativos Principais', 2)
        tabela = doc.add_table(rows=len(disc['principais'])+1, cols=3)
        tabela.style = 'Light Grid Accent 1'
        adicionar_borda_tabela(tabela)
        
        # Cabeçalho
        tabela.rows[0].cells[0].text = 'Descrição'
        tabela.rows[0].cells[1].text = 'UN'
        tabela.rows[0].cells[2].text = 'QTD'
        for cell in tabela.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Dados
        for i, (desc, un, qtd) in enumerate(disc['principais'], start=1):
            tabela.rows[i].cells[0].text = desc
            tabela.rows[i].cells[1].text = un
            tabela.rows[i].cells[2].text = formatar_numero(qtd, 2 if isinstance(qtd, float) else 0)
            tabela.rows[i].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Observações
        if disc['obs']:
            doc.add_paragraph()
            doc.add_heading('Observações', 2)
            doc.add_paragraph(f"⚠️ {disc['obs']}", style='List Bullet')
        
        doc.add_page_break()
    
    # =============================
    # APÊNDICES
    # =============================
    doc.add_heading('Apêndices', 1)
    
    doc.add_heading('Fontes de Dados', 2)
    p = doc.add_paragraph("Arquivos IFC, DXF e DWG processados:\n", style='List Bullet')
    p.add_run("• Estrutura: IFC R26\n")
    p.add_run("• Arquitetura: IFC R07 + R08\n")
    p.add_run("• Hidráulico: IFC rev.01 (348 - H00)\n")
    p.add_run("• Sanitário: IFC rev.01 (10 arquivos)\n")
    p.add_run("• PCI Civil: IFC rev.01 (4 arquivos)\n")
    p.add_run("• PCI Elétrico: IFC rev.01 (9 arquivos)\n")
    p.add_run("• Elétrico: IFC rev.01 (9 arquivos)\n")
    p.add_run("• Telefônico: IFC rev.01 (9 arquivos)\n")
    p.add_run("• SPDA: DWG (10 arquivos) — estimativa\n")
    p.add_run("• Ventilação: DXF R05 (30 MB)\n")
    p.add_run("• Exaustão: DXF R00 (18 MB)\n")
    p.add_run("• Ar-Condicionado: DXF R05 (39,9 MB)\n")
    p.add_run("• Alvenaria: DXF (5 pavimentos processados)\n")
    p.add_run("• Esquadrias: IFC R07 + JSON")
    
    doc.add_paragraph()
    
    doc.add_heading('Limitações Conhecidas', 2)
    limitacoes = [
        "Alvenaria: 4 pavimentos de garagem não processados (arquivos >300MB)",
        "Sanitário: Metragens de tubulação com valores zerados em 90% dos pavimentos",
        "PCI Civil: Metragem de tubulação subestimada (67m para 34 pavimentos)",
        "Elétrico: Diâmetros de eletrodutos e bitolas de cabos não especificados",
        "Telefônico: Cabos UTP, racks e patch panels não modelados",
        "SPDA: Quantitativos baseados em premissas técnicas — validar memorial",
        "Ventilação: Vazões e pressões de ventiladores pendentes de memorial",
        "Ar-Condicionado: Tubulações estimadas — extração de layers específicos pendente"
    ]
    
    for lim in limitacoes:
        doc.add_paragraph(f"• {lim}", style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Próximos Passos', 2)
    proximos = [
        "Processar 4 pavimentos de garagem em Alvenaria (G1, G2, G3, G5)",
        "Recalcular metragens de tubulação sanitária via geometria IFC ou DWG",
        "Validar metragem de tubulação PCI Civil com pranchas DWG",
        "Extrair diâmetros de eletrodutos e bitolas de cabos das legendas DWG",
        "Solicitar memorial descritivo de Ventilação (vazões, pressões)",
        "Solicitar memorial descritivo de SPDA (validar quantitativos estimados)",
        "Extrair layers de tubulação frigorifica em Ar-Condicionado",
        "Gerar planilhas executivas Excel compatíveis com Memorial Cartesiano"
    ]
    
    for prox in proximos:
        doc.add_paragraph(f"• {prox}", style='List Bullet')
    
    # Salvar documento
    output_path = "output/Briefing-Thozen-Electra-Completo.docx"
    os.makedirs("output", exist_ok=True)
    doc.save(output_path)
    
    file_size = os.path.getsize(output_path) / 1024  # KB
    print(f"✅ Briefing Word gerado: {output_path} ({file_size:.1f} KB)")
    return output_path

# ===========================
# GERAÇÃO DO EXCEL
# ===========================

def gerar_planilha_excel():
    """Gera planilha detalhada em Excel"""
    print("\n🔄 Gerando planilha Excel...")
    
    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # Remove planilha padrão
    
    # Estilos
    cabecalho_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    cabecalho_font = Font(color="FFFFFF", bold=True, size=11)
    total_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
    total_font = Font(bold=True, size=11)
    borda = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # =============================
    # ABA 1: RESUMO
    # =============================
    ws_resumo = wb.create_sheet("RESUMO", 0)
    
    # Título
    ws_resumo['A1'] = "RESUMO GERAL - THOZEN ELECTRA"
    ws_resumo['A1'].font = Font(size=16, bold=True, color="003366")
    ws_resumo.merge_cells('A1:F1')
    
    ws_resumo['A2'] = f"Data: {PROJETO['data']} | Cartesian Engenharia"
    ws_resumo.merge_cells('A2:F2')
    
    # Tabela consolidada
    ws_resumo['A4'] = "Disciplina"
    ws_resumo['B4'] = "Status"
    ws_resumo['C4'] = "Qtd Itens"
    ws_resumo['D4'] = "Custo Estimado"
    ws_resumo['E4'] = "Incerteza"
    ws_resumo['F4'] = "% do Total"
    
    for col in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws_resumo[f'{col}4'].fill = cabecalho_fill
        ws_resumo[f'{col}4'].font = cabecalho_font
        ws_resumo[f'{col}4'].alignment = Alignment(horizontal='center', vertical='center')
        ws_resumo[f'{col}4'].border = borda
    
    # Dados
    row = 5
    for disc in DISCIPLINAS:
        ws_resumo[f'A{row}'] = f"{disc['codigo']} - {disc['nome']}"
        ws_resumo[f'B{row}'] = disc['status']
        ws_resumo[f'C{row}'] = disc['qtd_itens']
        ws_resumo[f'D{row}'] = disc['custo_estimado']
        ws_resumo[f'D{row}'].number_format = 'R$ #,##0.00'
        ws_resumo[f'E{row}'] = disc['incerteza']
        ws_resumo[f'F{row}'] = disc['custo_estimado'] / TOTAL_CUSTO
        ws_resumo[f'F{row}'].number_format = '0.0%'
        
        for col in ['A', 'B', 'C', 'D', 'E', 'F']:
            ws_resumo[f'{col}{row}'].border = borda
        
        row += 1
    
    # Totais
    ws_resumo[f'A{row}'] = "TOTAL"
    ws_resumo[f'C{row}'] = TOTAL_ITENS
    ws_resumo[f'D{row}'] = TOTAL_CUSTO
    ws_resumo[f'D{row}'].number_format = 'R$ #,##0.00'
    ws_resumo[f'F{row}'] = 1.0
    ws_resumo[f'F{row}'].number_format = '0.0%'
    
    for col in ['A', 'B', 'C', 'D', 'E', 'F']:
        ws_resumo[f'{col}{row}'].fill = total_fill
        ws_resumo[f'{col}{row}'].font = total_font
        ws_resumo[f'{col}{row}'].border = borda
    
    # Larguras de coluna
    ws_resumo.column_dimensions['A'].width = 35
    ws_resumo.column_dimensions['B'].width = 15
    ws_resumo.column_dimensions['C'].width = 12
    ws_resumo.column_dimensions['D'].width = 18
    ws_resumo.column_dimensions['E'].width = 12
    ws_resumo.column_dimensions['F'].width = 12
    
    # Gráfico de pizza
    chart = PieChart()
    labels = Reference(ws_resumo, min_col=1, min_row=5, max_row=row-1)
    data = Reference(ws_resumo, min_col=4, min_row=4, max_row=row-1)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(labels)
    chart.title = "Distribuição de Custos por Disciplina"
    chart.height = 12
    chart.width = 20
    ws_resumo.add_chart(chart, "A" + str(row + 3))
    
    # Congelar painéis
    ws_resumo.freeze_panes = 'A5'
    
    # =============================
    # ABAS POR DISCIPLINA
    # =============================
    for disc in DISCIPLINAS:
        nome_aba = f"{disc['codigo']}-{disc['nome'][:20]}"
        ws = wb.create_sheet(nome_aba)
        
        # Título
        ws['A1'] = f"{disc['codigo']} - {disc['nome']}"
        ws['A1'].font = Font(size=14, bold=True, color="003366")
        ws.merge_cells('A1:G1')
        
        ws['A2'] = f"Status: {disc['status']} | Fonte: {disc['fonte']}"
        ws.merge_cells('A2:G2')
        
        # Cabeçalho
        cabecalhos = ['Código', 'Descrição', 'UN', 'QTD', 'Preço Unit.', 'Total', 'Observação']
        for col_idx, cabecalho in enumerate(cabecalhos, start=1):
            col_letter = get_column_letter(col_idx)
            ws[f'{col_letter}4'] = cabecalho
            ws[f'{col_letter}4'].fill = cabecalho_fill
            ws[f'{col_letter}4'].font = cabecalho_font
            ws[f'{col_letter}4'].alignment = Alignment(horizontal='center', vertical='center')
            ws[f'{col_letter}4'].border = borda
        
        # Dados
        row = 5
        for item_idx, (desc, un, qtd) in enumerate(disc['principais'], start=1):
            ws[f'A{row}'] = f"{disc['codigo']}.{item_idx:03d}"
            ws[f'B{row}'] = desc
            ws[f'C{row}'] = un
            ws[f'D{row}'] = qtd
            ws[f'D{row}'].number_format = '#,##0.00' if isinstance(qtd, float) else '#,##0'
            ws[f'E{row}'] = ""  # Preço unitário (não temos)
            ws[f'F{row}'] = ""  # Total (não temos)
            ws[f'G{row}'] = disc['obs'] if row == 5 else ""
            
            for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
                ws[f'{col}{row}'].border = borda
            
            row += 1
        
        # Rodapé
        ws[f'A{row+1}'] = f"Fonte de dados: {disc['fonte']}"
        ws[f'A{row+1}'].font = Font(italic=True, size=9)
        
        # Larguras de coluna
        ws.column_dimensions['A'].width = 12
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 8
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 15
        ws.column_dimensions['G'].width = 50
        
        # Congelar painéis
        ws.freeze_panes = 'A5'
        
        # Filtros automáticos
        ws.auto_filter.ref = f"A4:G{row-1}"
    
    # Salvar planilha
    output_path = "output/Quantitativos-Thozen-Electra-Detalhados.xlsx"
    wb.save(output_path)
    
    file_size = os.path.getsize(output_path) / 1024  # KB
    print(f"✅ Planilha Excel gerada: {output_path} ({file_size:.1f} KB)")
    return output_path

# ===========================
# MAIN
# ===========================

if __name__ == "__main__":
    print("=" * 60)
    print("GERAÇÃO DE BRIEFING E PLANILHA - THOZEN ELECTRA")
    print("=" * 60)
    
    try:
        # Gerar Word
        word_path = gerar_briefing_word()
        
        # Gerar Excel
        excel_path = gerar_planilha_excel()
        
        print("\n" + "=" * 60)
        print("✅ CONCLUÍDO COM SUCESSO")
        print("=" * 60)
        print(f"\n📄 Briefing Word: {word_path}")
        print(f"📊 Planilha Excel: {excel_path}")
        print(f"\n📁 Localização: {os.path.abspath('output/')}")
        print("\n" + "=" * 60)
        
    except Exception as e:
        print(f"\n❌ ERRO: {e}")
        import traceback
        traceback.print_exc()
