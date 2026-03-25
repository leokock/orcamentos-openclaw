#!/usr/bin/env python3
"""
Gera doc Word do Electra R02 com rastreabilidade item a item.
Cada item mostra de onde veio (IFC/DXF, Referencia, Estimado).
"""
import json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import openpyxl

# === CONFIG ===
XLSX = os.path.expanduser("~/orcamentos/executivos/thozen-electra/entregas/CTN-TZN_ELT-Orcamento-Executivo-R02-FORMATO-EII.xlsx")
OUTPUT = os.path.expanduser("~/orcamentos/executivos/thozen-electra/entregas/CTN-TZN_ELT-Orcamento-Executivo-R02-DOC-RASTREAVEL.docx")

AC = 36088.85
CUB = 3019.26
TOTAL_GERAL = 169522402.07

# Colors
DARK_BLUE = RGBColor(0x2C, 0x3E, 0x50)
ACCENT = RGBColor(0x29, 0x80, 0xB9)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

COLOR_MAP = {
    '00C6EFCE': ('IFC/DXF', 'C6EFCE'),
    '00FFEB9C': ('Referencia', 'FFEB9C'),
    '00FFC7CE': ('Estimado', 'FFC7CE'),
}

GREEN_HEX = "C6EFCE"
YELLOW_HEX = "FFEB9C"
RED_HEX = "FFC7CE"
GRAY_HEX = "F0F0F0"
WHITE_HEX = "FFFFFF"
HEADER_HEX = "2C3E50"

# === HELPERS ===
def set_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def brl(v):
    if v is None: return "-"
    return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def brl_int(v):
    if v is None: return "-"
    return f"R$ {v:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")

def get_source(cell):
    fill = cell.fill
    for attr in ['fgColor', 'start_color']:
        c = getattr(fill, attr, None)
        if c and c.rgb and c.rgb in COLOR_MAP:
            return COLOR_MAP[c.rgb]
    return ('Sem tag', GRAY_HEX)

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE if level == 1 else ACCENT
    return h

def add_header_row(table, cols, widths=None):
    row = table.rows[0] if table.rows else table.add_row()
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        run.font.name = "Arial"
        set_shading(cell, HEADER_HEX)

def add_item_row(table, cells_data, source_tag, source_color):
    row = table.add_row()
    for i, (text, align) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(7.5)
        run.font.name = "Arial"
    # Last cell = source with color
    src_cell = row.cells[len(cells_data)]
    src_cell.text = ""
    p = src_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(source_tag)
    run.font.size = Pt(7)
    run.font.name = "Arial"
    run.font.bold = True
    set_shading(src_cell, source_color)

# === EXTRACT DATA FROM XLSX ===
print("Lendo xlsx...")
wb = openpyxl.load_workbook(XLSX)

DETAIL_TABS = [
    'ELETRICO', 'HIDROSSANITARIO', 'PPCI', 'TELECOMUNICACAO', 'CLIMATIZACAO',
    'ESQUADRIAS', 'DRYWALL', 'ILUMINACAO', 'GAS', 'LOUCAS E METAIS',
    'PISCINA', 'Equipamentos Especiais', 'IMPERMEABILIZACAO', 'MOBILIARIO',
    'Bombeamento - Extra', 'AUTOMACAO'
]

# Map accented names to sheet names
SHEET_MAP = {}
for sn in wb.sheetnames:
    import unicodedata
    norm = unicodedata.normalize('NFD', sn).encode('ascii', 'ignore').decode().upper()
    norm = norm.replace(' ', '').replace('-','').replace('|','')
    SHEET_MAP[norm] = sn

def find_sheet(name):
    import unicodedata
    norm = unicodedata.normalize('NFD', name).encode('ascii', 'ignore').decode().upper()
    norm = norm.replace(' ', '').replace('-','').replace('|','')
    if norm in SHEET_MAP:
        return SHEET_MAP[norm]
    # Fuzzy
    for k, v in SHEET_MAP.items():
        if norm[:8] in k:
            return v
    return None

all_tabs = {}
for tab_key in DETAIL_TABS:
    sheet_name = find_sheet(tab_key)
    if not sheet_name:
        print(f"  SKIP: {tab_key} not found")
        continue
    ws = wb[sheet_name]

    # Get origem
    origem = ""
    for r in range(1, 5):
        for c in range(1, 10):
            v = ws.cell(r, c).value
            if v and isinstance(v, str) and 'Origem' in v:
                for cc in range(c+1, c+3):
                    ov = ws.cell(r, cc).value
                    if ov:
                        origem = str(ov)
                        break

    # Get revisao
    revisao = ""
    for r in range(1, 5):
        for c in range(1, 10):
            v = ws.cell(r, c).value
            if v and isinstance(v, str) and ('Revis' in v or 'revis' in v):
                for cc in range(c+1, c+3):
                    rv = ws.cell(r, cc).value
                    if rv:
                        revisao = str(rv)
                        break

    # Projetistas do Electra por disciplina
    PROJETISTAS = {
        'ELÉTRICO': 'Franzmann',
        'HIDROSSANITÁRIO': 'Franzmann',
        'PPCI': 'Franzmann',
        'TELECOMUNICAÇÃO': 'Franzmann',
        'ILUMINAÇÃO': 'Franzmann',
        'GÁS': 'Franzmann',
        'CLIMATIZAÇÃO': 'Value',
        'ESQUADRIAS': 'WM',
        'IMPERMEABILIZAÇÃO': 'W.Thomaz',
        'DRYWALL': '',
        'LOUÇAS E METAIS': '',
        'PISCINA': '',
        'Equipamentos Especiais': '',
        'MOBILIÁRIO': '',
        'Bombeamento - Extra': '',
        'AUTOMAÇÃO': '',
    }
    projetista = PROJETISTAS.get(sheet_name, '')

    # Build specific source tags per tab — referencing Electra project files
    def get_specific_source(cell, desc_str):
        generic_tag, color = get_source(cell)
        if color == GREEN_HEX:
            # Green = extracted from Electra project files
            proj_prefix = f"Proj. {projetista} " if projetista else "Proj. Electra "
            if '— IFC' in desc_str or '— DXF' not in desc_str:
                # Default to IFC for green items
                for part in origem.split('+'):
                    part = part.strip()
                    if 'IFC' in part and ('rev' in part.lower() or 'R' in part):
                        return f"{proj_prefix}{part}".strip()[:28], color
                if 'IFC' in origem:
                    return f"{proj_prefix}IFC".strip()[:28], color
            if '— DXF' in desc_str:
                for part in origem.split('+'):
                    part = part.strip()
                    if 'DXF' in part:
                        return f"{proj_prefix}{part}".strip()[:28], color
                return f"{proj_prefix}DXF".strip()[:28], color
            if '— BIM' in desc_str:
                return f"{proj_prefix}IFC (BIM)".strip()[:28], color
            # Fallback
            return f"{proj_prefix}IFC/DXF".strip()[:28], color
        elif color == YELLOW_HEX:
            # Yellow = parametric reference — NEVER name other clients' projects
            dl = desc_str.lower()
            if 'sinapi' in dl:
                return 'SINAPI', color
            elif 'base' in dl and 'R$' in desc_str:
                return 'PU base Cartesian', color
            elif 'estimado' in dl or 'indice' in dl:
                return 'Indice param.', color
            elif projetista and ('briefing' in origem.lower() or 'briefing' in dl):
                return f'Briefing {projetista}', color
            else:
                return 'Param. base Cartesian', color
        elif color == RED_HEX:
            return 'Estimado (s/ projeto)', color
        else:
            return generic_tag, color

    items = []
    for r in range(7, ws.max_row + 1):
        desc = ws.cell(r, 5).value
        total = ws.cell(r, 12).value
        if not desc or not total or not isinstance(total, (int, float)):
            continue
        desc_str = str(desc)
        if any(x in desc_str.upper() for x in ['R$/M', 'LEGENDA', 'TOTAL GERAL']):
            continue

        qty = ws.cell(r, 9).value or ws.cell(r, 6).value
        unit = ws.cell(r, 10).value
        pu = ws.cell(r, 11).value
        source_tag, source_color = get_specific_source(ws.cell(r, 5), desc_str)
        pav = ws.cell(r, 2).value or ''

        items.append({
            'desc': desc_str[:70],
            'pav': str(pav)[:25] if pav else '',
            'qty': float(qty) if qty and isinstance(qty, (int, float)) else None,
            'unit': str(unit) if unit else '',
            'pu': float(pu) if pu and isinstance(pu, (int, float)) else None,
            'total': float(total),
            'source_tag': source_tag,
            'source_color': source_color,
        })

    # Sanitize origem — never reference other clients' projects by name
    origem_clean = origem.replace('Elizabeth II', 'proj. referencia similar').replace('elizabeth II', 'proj. referencia similar')
    origem_clean = origem_clean.replace('Ref. proj. referencia similar', 'Ref. projeto similar')

    tab_total = sum(it['total'] for it in items)
    all_tabs[sheet_name] = {'origem': origem_clean, 'items': items, 'total': tab_total}
    print(f"  {sheet_name}: {len(items)} itens | R$ {tab_total:,.0f}")

# Ger_Tec e Adm
ws_ger = wb['Ger_Tec e Adm']
ger_items = []
for r in range(9, ws_ger.max_row + 1):
    nivel = ws_ger.cell(r, 8).value
    desc = ws_ger.cell(r, 10).value
    unit = ws_ger.cell(r, 11).value
    qty = ws_ger.cell(r, 12).value
    pu = ws_ger.cell(r, 13).value
    total = ws_ger.cell(r, 14).value
    obs = str(ws_ger.cell(r, 15).value or '')

    if not desc or not total or not isinstance(total, (int, float)):
        continue
    if nivel in ['CELULA CONSTRUTIVA', 'ETAPA', 'CÉLULA CONSTRUTIVA']:
        continue

    obs_lower = obs.lower()
    if 'contrato' in obs_lower:
        # Extract contractor name from obs
        tag = 'Contrato'
        for name in ['Battisti', 'Favola', 'DMA', 'Zeplin', 'Liberte', 'franzmann', 'value', 'W.Thomaz', 'WM', 'otus', 'triunfo']:
            if name.lower() in obs_lower:
                tag = f'Contrato {name.title()}'
                break
        source_tag, source_color = tag, GREEN_HEX
    elif 'repessado' in obs_lower or 'gessele' in obs_lower:
        source_tag, source_color = 'Thozen 05/03/26', GREEN_HEX
    elif 'estimado' in obs_lower:
        source_tag, source_color = 'Estimado (s/ proj)', RED_HEX
    elif 'sinapi' in obs_lower:
        source_tag, source_color = 'SINAPI', YELLOW_HEX
    elif 'mediana' in obs_lower or 'parametr' in obs_lower:
        source_tag, source_color = 'Param. base Cartesian', YELLOW_HEX
    else:
        source_tag, source_color = 'Param. base Cartesian', YELLOW_HEX

    ger_items.append({
        'desc': str(desc)[:70],
        'pav': str(nivel)[:25] if nivel else '',
        'qty': float(qty) if qty and isinstance(qty, (int, float)) else None,
        'unit': str(unit) if unit else '',
        'pu': float(pu) if pu and isinstance(pu, (int, float)) else None,
        'total': float(total),
        'source_tag': source_tag,
        'source_color': source_color,
        'obs': obs[:50]
    })

ger_total = sum(it['total'] for it in ger_items)
print(f"  Ger_Tec e Adm: {len(ger_items)} itens | R$ {ger_total:,.0f}")

# Resumo Estrutura
ws_est = wb['Resumo Estrutura']
est_items = []
for r in range(4, ws_est.max_row + 1):
    disc = ws_est.cell(r, 2).value
    pav = ws_est.cell(r, 3).value
    etapa = ws_est.cell(r, 4).value
    proj = ws_est.cell(r, 5).value
    concreto = ws_est.cell(r, 8).value
    forma = ws_est.cell(r, 9).value
    if not disc:
        continue
    est_items.append({
        'disc': str(disc),
        'pav': str(pav) if pav else '',
        'etapa': str(etapa) if etapa else '',
        'proj': str(proj) if proj else '',
        'concreto': str(concreto) if concreto else '',
        'forma': str(forma) if forma else '',
    })
print(f"  Resumo Estrutura: {len(est_items)} itens")

# === EAP Summary ===
ws_eap = wb['EAP Análise Revisão']
eap_rows = []
for r in range(3, 29):
    etapa = ws_eap.cell(r, 1).value
    valor = ws_eap.cell(r, 2).value
    pct = ws_eap.cell(r, 3).value
    rsm2 = ws_eap.cell(r, 4).value
    med = ws_eap.cell(r, 5).value
    diff = ws_eap.cell(r, 8).value
    obs = ws_eap.cell(r, 9).value
    if not etapa or not valor:
        continue
    eap_rows.append({
        'etapa': str(etapa),
        'valor': float(valor),
        'pct': float(pct) if pct else 0,
        'rsm2': float(rsm2) if rsm2 else 0,
        'has_benchmark': med is not None,
        'diff': float(diff) if diff and isinstance(diff, (int, float)) else None,
        'obs': str(obs) if obs else '',
    })

# === GENERATE DOC ===
print("\nGerando documento...")
doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.space_before = Pt(1)

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# === CAPA ===
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CARTESIAN ENGENHARIA")
run.font.size = Pt(11)
run.font.color.rgb = DARK_BLUE
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ORCAMENTO EXECUTIVO")
run.font.size = Pt(22)
run.font.color.rgb = DARK_BLUE
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ELECTRA TOWERS")
run.font.size = Pt(18)
run.font.color.rgb = ACCENT
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Revisao R02 | Formato EII | Rastreabilidade Completa")
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_paragraph()

# Dados do projeto
dados = [
    ("Empresa", "Thozen"), ("Cidade/UF", "Porto Belo, SC"),
    ("AC", "36.088,85 m2"), ("Pavimentos", "35"), ("UR", "342 + 6 com."),
    ("Vagas", "305"), ("Subsolos", "1"), ("Prazo", "36 meses"),
    ("CUB/SC fev/2026", "R$ 3.019,26"), ("CUB Ratio", "1,56"),
]
table = doc.add_table(rows=0, cols=4)
table.style = 'Table Grid'
for i in range(0, len(dados), 2):
    row = table.add_row()
    for j in range(2):
        if i + j < len(dados):
            k, v = dados[i + j]
            row.cells[j*2].text = k
            row.cells[j*2].paragraphs[0].runs[0].font.bold = True if row.cells[j*2].paragraphs[0].runs else None
            row.cells[j*2+1].text = v
            for c in [row.cells[j*2], row.cells[j*2+1]]:
                for run in c.paragraphs[0].runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
    if i % 4 == 0:
        for c in row.cells:
            set_shading(c, GRAY_HEX)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Marco 2026")
run.font.color.rgb = GRAY

doc.add_page_break()

# === 1. LEGENDA DE RASTREABILIDADE ===
add_styled_heading(doc, "1. Legenda de Rastreabilidade", level=1)

p = doc.add_paragraph()
p.add_run("Cada item do orcamento e classificado pela origem dos seus dados:").font.color.rgb = GRAY

table_leg = doc.add_table(rows=4, cols=3)
table_leg.style = 'Table Grid'
add_header_row(table_leg, ["Cor", "Classificacao", "Significado"])
leg_data = [
    (GREEN_HEX, "Projeto Electra / Confirmado", "Extraido de IFC/DXF dos projetistas do Electra ou confirmado pela Thozen"),
    (YELLOW_HEX, "Parametrico / Referencia", "PU ou qtd baseado na base Cartesian (75 executivos) ou SINAPI"),
    (RED_HEX, "Estimado (sem projeto)", "Sem fonte confiavel - requer cotacao ou projeto especifico para proxima revisao"),
]
for i, (color, tag, desc) in enumerate(leg_data):
    row = table_leg.rows[i + 1]
    set_shading(row.cells[0], color)
    row.cells[0].text = ""
    row.cells[1].text = tag
    row.cells[2].text = desc
    for c in row.cells:
        for run in c.paragraphs[0].runs:
            run.font.size = Pt(8.5)
            run.font.name = 'Arial'

doc.add_paragraph()

# Maturidade summary
verde_total = sum(
    sum(it['total'] for it in tab['items'] if it['source_color'] == GREEN_HEX)
    for tab in all_tabs.values()
) + sum(it['total'] for it in ger_items if it['source_color'] == GREEN_HEX)

amarelo_total = sum(
    sum(it['total'] for it in tab['items'] if it['source_color'] == YELLOW_HEX)
    for tab in all_tabs.values()
) + sum(it['total'] for it in ger_items if it['source_color'] == YELLOW_HEX)

vermelho_total = sum(
    sum(it['total'] for it in tab['items'] if it['source_color'] == RED_HEX)
    for tab in all_tabs.values()
) + sum(it['total'] for it in ger_items if it['source_color'] == RED_HEX)

total_tagged = verde_total + amarelo_total + vermelho_total

p = doc.add_paragraph()
run = p.add_run("Distribuicao: ")
run.font.bold = True
p.add_run(f"IFC/DXF {verde_total/TOTAL_GERAL*100:.1f}% | Referencia {amarelo_total/TOTAL_GERAL*100:.1f}% | Estimado {vermelho_total/TOTAL_GERAL*100:.1f}%")

doc.add_paragraph()

# === 2. EAP RESUMO ===
add_styled_heading(doc, "2. EAP Resumo com Origem", level=1)

table_eap = doc.add_table(rows=1, cols=5)
table_eap.style = 'Table Grid'
add_header_row(table_eap, ["Etapa", "Valor (R$)", "R$/m2", "%", "Origem"])

for e in eap_rows:
    row = table_eap.add_row()
    row.cells[0].text = e['etapa'][:45]
    row.cells[1].text = brl_int(e['valor'])
    row.cells[2].text = f"{e['rsm2']:.0f}"
    row.cells[3].text = f"{e['pct']*100:.1f}%"
    row.cells[4].text = e['obs'][:40]
    for c in row.cells:
        for run in c.paragraphs[0].runs:
            run.font.size = Pt(7.5)
            run.font.name = 'Arial'
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Total row
row = table_eap.add_row()
row.cells[0].text = "TOTAL"
row.cells[1].text = brl_int(TOTAL_GERAL)
row.cells[2].text = f"{TOTAL_GERAL/AC:.0f}"
row.cells[3].text = "100%"
for c in row.cells:
    set_shading(c, HEADER_HEX)
    for run in c.paragraphs[0].runs:
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.name = 'Arial'

doc.add_page_break()

# === 3. ESTRUTURA (RESUMO) ===
add_styled_heading(doc, "3. Estrutura - Quantitativos IFC R26", level=1)

p = doc.add_paragraph()
p.add_run("Fonte: ").font.bold = True
p.add_run("Modelo IFC R26 (DMA + Zeplin + Liberte). Todos os quantitativos abaixo sao IFC/DXF (verde).")

table_est = doc.add_table(rows=1, cols=6)
table_est.style = 'Table Grid'
add_header_row(table_est, ["Disciplina", "Pavimento", "Etapa", "Projetos", "Concreto (m3)", "Forma (m2)"])

for it in est_items:
    row = table_est.add_row()
    row.cells[0].text = it['disc']
    row.cells[1].text = it['pav']
    row.cells[2].text = it['etapa']
    row.cells[3].text = it['proj']
    row.cells[4].text = it['concreto']
    row.cells[5].text = it['forma']
    for c in row.cells:
        set_shading(c, GREEN_HEX)
        for run in c.paragraphs[0].runs:
            run.font.size = Pt(8)
            run.font.name = 'Arial'

doc.add_paragraph()

# === 4. GERENCIAMENTO ===
add_styled_heading(doc, "4. Gerenciamento Tecnico e Administrativo", level=1)

p = doc.add_paragraph()
p.add_run("Origem: ").font.bold = True
p.add_run("Valores confirmados pela Thozen na reuniao 05/03/2026 (verde), com contratos de projetistas identificados. Itens sem confirmacao sao parametricos (amarelo) ou estimados (vermelho).")

table_ger = doc.add_table(rows=1, cols=6)
table_ger.style = 'Table Grid'
add_header_row(table_ger, ["Descricao", "Qtd", "Unid", "PU (R$)", "Total (R$)", "Fonte"])

for it in ger_items:
    row = table_ger.add_row()
    cells_data = [
        (it['desc'], WD_ALIGN_PARAGRAPH.LEFT),
        (f"{it['qty']:,.1f}" if it['qty'] else "-", WD_ALIGN_PARAGRAPH.CENTER),
        (it['unit'], WD_ALIGN_PARAGRAPH.CENTER),
        (brl(it['pu']) if it['pu'] else "-", WD_ALIGN_PARAGRAPH.RIGHT),
        (brl_int(it['total']), WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    for i, (text, align) in enumerate(cells_data):
        row.cells[i].text = str(text)
        row.cells[i].paragraphs[0].alignment = align
        for run in row.cells[i].paragraphs[0].runs:
            run.font.size = Pt(7.5)
            run.font.name = 'Arial'
    # Source cell
    row.cells[5].text = it['source_tag']
    row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_shading(row.cells[5], it['source_color'])
    for run in row.cells[5].paragraphs[0].runs:
        run.font.size = Pt(7)
        run.font.name = 'Arial'
        run.font.bold = True

doc.add_page_break()

# === 5+. DISCIPLINAS DETALHADAS ===
section_num = 5
for sheet_name, tab_data in all_tabs.items():
    items = tab_data['items']
    if not items:
        continue

    add_styled_heading(doc, f"{section_num}. {sheet_name}", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Origem: ")
    run.font.bold = True
    p.add_run(tab_data['origem'] or "Ver legenda")

    p = doc.add_paragraph()
    p.add_run(f"Total: {brl_int(tab_data['total'])} | {len(items)} itens | R$/m2: {tab_data['total']/AC:.2f}")

    # Count by source
    from collections import Counter
    src_count = Counter(it['source_tag'] for it in items)
    src_val = {}
    for it in items:
        src_val[it['source_tag']] = src_val.get(it['source_tag'], 0) + it['total']

    p = doc.add_paragraph()
    parts = []
    for tag, count in src_count.most_common():
        val = src_val[tag]
        parts.append(f"{tag}: {count} itens ({val/tab_data['total']*100:.0f}%)")
    p.add_run("  |  ".join(parts)).font.size = Pt(8.5)

    doc.add_paragraph()

    # Items table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    add_header_row(table, ["Descricao", "Qtd", "Unid", "PU (R$)", "Total (R$)", "Fonte"])

    for it in items:
        row = table.add_row()
        qty_str = f"{it['qty']:,.1f}" if it['qty'] else "-"
        pu_str = brl(it['pu']) if it['pu'] else "-"

        cells = [
            (it['desc'], WD_ALIGN_PARAGRAPH.LEFT),
            (qty_str, WD_ALIGN_PARAGRAPH.CENTER),
            (it['unit'], WD_ALIGN_PARAGRAPH.CENTER),
            (pu_str, WD_ALIGN_PARAGRAPH.RIGHT),
            (brl_int(it['total']), WD_ALIGN_PARAGRAPH.RIGHT),
        ]
        for i, (text, align) in enumerate(cells):
            row.cells[i].text = str(text)
            row.cells[i].paragraphs[0].alignment = align
            for run in row.cells[i].paragraphs[0].runs:
                run.font.size = Pt(7.5)
                run.font.name = 'Arial'

        # Source
        row.cells[5].text = it['source_tag']
        row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_shading(row.cells[5], it['source_color'])
        for run in row.cells[5].paragraphs[0].runs:
            run.font.size = Pt(7)
            run.font.name = 'Arial'
            run.font.bold = True

    section_num += 1
    if section_num % 4 == 0:
        doc.add_page_break()

# === PREMISSAS ===
doc.add_page_break()
add_styled_heading(doc, f"{section_num}. Premissas e Limitacoes", level=1)

premissas = [
    "CUB referencia: R$ 3.019,26 (SC, fevereiro/2026) - fonte Sinduscon-SC",
    "Area construida: 36.088,85 m2 (projeto arquitetonico)",
    "Disciplinas IFC/DXF (verde): Eletrico R02, Hidrossanitario R01, PPCI R01, Estrutura IFC R26, Esquadrias BIM",
    "Disciplinas parametricas (amarelo): mediana base Cartesian (75 executivos calibrados)",
    "Parametricos: mediana base Cartesian (75 exec) aplicada sobre projeto de porte e padrao similar",
    "Contencao estimada sem projeto (vermelho) - requer projeto especifico para R03",
    "Gerenciamento: valores confirmados pela Thozen na reuniao 05/03/2026, com contratos de projetistas",
    "PUs-chave: contrapiso R$ 12,10/m2, reboco R$ 7,00/m2, manta R$ 82,15/m2, reboco ext R$ 15,25/m2 (medianas base 37 exec)",
    "Valores nao incluem terreno, incorporacao, marketing ou despesas financeiras",
    "Proxima revisao (R03): processar DWGs climatizacao, detalhar MO eletrica e hidro, obter projeto contencao",
]

for item in premissas:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Cartesian Engenharia")
run.font.bold = True
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Integracao de escopo, custo e prazo")
run.font.color.rgb = GRAY
run.font.size = Pt(9)

# === SAVE ===
doc.save(OUTPUT)
print(f"\nGerado: {OUTPUT}")
print(f"Total de itens rastreados: {sum(len(t['items']) for t in all_tabs.values()) + len(ger_items) + len(est_items)}")
