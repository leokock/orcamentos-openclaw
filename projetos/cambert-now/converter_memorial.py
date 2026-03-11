#!/usr/bin/env python3.11
# -*- coding: utf-8 -*-
"""
Converte memorial-descritivo.md em documento Word profissional
Cartesian Engenharia - Março 2026
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_page_number(section):
    """Adiciona número de página ao rodapé"""
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # Criar parágrafo para o rodapé
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar texto "Cartesian Engenharia | Confidencial  "
    run = p.add_run("Cartesian Engenharia | Confidencial  ")
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    
    # Adicionar número de página
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)

def add_header(section, text):
    """Adiciona cabeçalho"""
    header = section.header
    header.is_linked_to_previous = False
    
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    run.font.italic = True

def configure_styles(doc):
    """Configura estilos do documento"""
    styles = doc.styles
    
    # Estilo H1
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)  # Azul escuro
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Estilo H2
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)  # Azul
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(5)
    
    # Estilo H3
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(0x44, 0x44, 0x44)  # Cinza escuro
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(4)
    
    # Estilo Normal (corpo)
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

def create_cover_page(doc):
    """Cria página de capa"""
    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(6)
    run = p.add_run('MEMORIAL DESCRITIVO — NOW RESIDENCE')
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('Edifício Residencial Multifamiliar')
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    
    # Endereço
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('Rua Luiz Berlim, 123, Centro, Itajaí/SC')
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    
    # Cartesian + Data
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('Cartesian Engenharia — Março 2026')
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Quebra de página após capa
    doc.add_page_break()

def process_markdown_line(doc, line, in_list=False):
    """Processa uma linha do markdown"""
    line = line.rstrip()
    
    # Headers H1 (##)
    if line.startswith('## ') and not line.startswith('###'):
        title = line[3:].strip()
        doc.add_heading(title, level=1)
        return 'h1'
    
    # Headers H2 (###)
    elif line.startswith('### ') and not line.startswith('####'):
        title = line[4:].strip()
        doc.add_heading(title, level=2)
        return 'h2'
    
    # Headers H3 (####)
    elif line.startswith('#### '):
        title = line[5:].strip()
        doc.add_heading(title, level=3)
        return 'h3'
    
    # Separador horizontal (---)
    elif line.strip() == '---':
        return 'separator'
    
    # Bullet points (- ou *)
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        p = doc.add_paragraph(style='List Bullet')
        add_formatted_text(p, text)
        return 'list'
    
    # Linhas vazias
    elif not line.strip():
        return 'empty'
    
    # Texto normal
    else:
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        return 'normal'

def add_formatted_text(paragraph, text):
    """Adiciona texto com formatação para [SUGERIDO] e **bold**"""
    # Processar [SUGERIDO]
    parts = re.split(r'(\[SUGERIDO\])', text)
    
    for part in parts:
        if part == '[SUGERIDO]':
            run = paragraph.add_run(part)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0xD6, 0x8C, 0x00)  # Laranja/amarelo escuro
        else:
            # Processar **bold** dentro do texto
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for bold_part in bold_parts:
                if bold_part.startswith('**') and bold_part.endswith('**'):
                    run = paragraph.add_run(bold_part[2:-2])
                    run.font.bold = True
                else:
                    paragraph.add_run(bold_part)

def convert_markdown_to_docx(md_file, docx_file):
    """Converte arquivo markdown para docx"""
    print("Iniciando conversão...")
    
    # Criar documento
    doc = Document()
    
    # Configurar margens (2.5 cm top/bottom, 2.0 cm left/right)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    
    # Configurar estilos
    configure_styles(doc)
    
    # Criar capa
    create_cover_page(doc)
    
    # Adicionar cabeçalho e rodapé (para todas as seções exceto a primeira)
    for i, section in enumerate(doc.sections):
        if i > 0:  # Pula a capa
            add_header(section, "Memorial Descritivo — NOW Residence")
            add_page_number(section)
    
    # Se só há uma seção (capa), criar nova seção para o conteúdo
    if len(doc.sections) == 1:
        doc.add_section()
        section = doc.sections[-1]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        add_header(section, "Memorial Descritivo — NOW Residence")
        add_page_number(section)
    
    # Ler e processar markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    previous_type = None
    skip_first_h1 = True  # Pular o primeiro # MEMORIAL DESCRITIVO
    
    for line in lines:
        # Pular primeira linha (título principal já está na capa)
        if skip_first_h1 and line.strip().startswith('# MEMORIAL DESCRITIVO'):
            skip_first_h1 = False
            continue
        
        # Pular segunda linha (subtítulo já está na capa)
        if line.strip().startswith('## NOW RESIDENCE'):
            continue
        
        line_type = process_markdown_line(doc, line)
        
        # Adicionar quebra de página antes de H1 (seções principais)
        if line_type == 'h1' and previous_type not in [None, 'h1', 'separator']:
            doc.add_page_break()
        
        previous_type = line_type
    
    # Salvar documento
    doc.save(docx_file)
    print(f"Documento salvo em: {docx_file}")

if __name__ == '__main__':
    import os
    
    # Caminhos dos arquivos
    base_dir = os.path.expanduser('~/orcamentos/projetos/cambert-now')
    md_file = os.path.join(base_dir, 'memorial-descritivo.md')
    docx_file = os.path.join(base_dir, 'NOW-Residence-Memorial-Descritivo.docx')
    
    # Converter
    convert_markdown_to_docx(md_file, docx_file)
    print("✅ Conversão concluída com sucesso!")
